from __future__ import annotations

import json
import math
from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables"
ASSETS = ROOT / "assets"
TABLES = ROOT / "tables"
OUT.mkdir(parents=True, exist_ok=True)
S = json.loads((ROOT / "analysis_summary.json").read_text(encoding="utf-8"))
DATE = "14 August 2026"
SOURCE_PAGE = "https://data.cdc.gov/National-Center-for-Health-Statistics/Provisional-COVID-19-Deaths-by-Sex-and-Age/9bhg-hcku"
SOURCE_CSV = "https://data.cdc.gov/api/views/9bhg-hcku/rows.csv?accessType=DOWNLOAD"
CDC_WONDER = "https://wonder.cdc.gov/mcd-icd10-provisional.html"

NAVY = "16324F"
BLUE = "2E75B6"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "666666"
DARK = "222222"
RED = "B5403A"
GREEN = "2F6B4F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color="B7C9D6", size="8"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_field(run, instruction):
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instruction
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def configure_document(doc: Document, week: int):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.70)
    sec.left_margin = Inches(0.78)
    sec.right_margin = Inches(0.78)
    sec.header_distance = Inches(0.28)
    sec.footer_distance = Inches(0.30)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 12),
        ("Heading 1", 18, NAVY, 14, 7),
        ("Heading 2", 14, BLUE, 11, 5),
        ("Heading 3", 11.5, GREEN, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), st.font.name)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Code Block" not in [s.name for s in styles]:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code.font.size = Pt(7.5)
        code.font.color.rgb = RGBColor.from_string("1F1F1F")
        code.paragraph_format.space_after = Pt(0)
        code.paragraph_format.line_spacing = 1.0
    if "Caption Custom" not in [s.name for s in styles]:
        cap = styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Aptos"
        cap.font.size = Pt(8.8)
        cap.font.italic = True
        cap.font.color.rgb = RGBColor.from_string(MID_GRAY)
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(8)

    header = sec.header
    p = header.paragraphs[0]
    p.text = f"DATA SCIENCE WITH PYTHON  |  WEEK {week}  |  CDC COVID-19 MORTALITY DATA"
    p.style = styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(BLUE)
    add_bottom_border(p, color="B7C9D6", size="6")

    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Page ")
    r.font.size = Pt(8)
    add_field(p.add_run(), "PAGE")
    p.add_run("  |  Prepared 14 August 2026")
    for r in p.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MID_GRAY)

    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_title_page(doc, title, week, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("DATA SCIENCE WITH PYTHON")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(27)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)

    line = doc.add_paragraph("━" * 47)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = ["WEEK", "DATE", "DATASET"]
    values = [str(week), DATE, "Provisional COVID-19 Deaths by Sex and Age"]
    for i, (label, value) in enumerate(zip(labels, values)):
        table.columns[0].width = Inches(1.35)
        table.columns[1].width = Inches(4.8)
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        set_cell_shading(table.cell(i, 0), NAVY)
        set_cell_shading(table.cell(i, 1), "F4F8FB")
        for r in table.cell(i, 0).paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.bold = True
            r.font.size = Pt(9)
        for r in table.cell(i, 1).paragraphs[0].runs:
            r.font.size = Pt(10)
        for c in table.rows[i].cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c, 110, 130, 110, 130)

    doc.add_paragraph("")
    identity = doc.add_table(rows=4, cols=2)
    identity.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, label in enumerate(["Student name", "Student ID", "Institution", "Mentor / supervisor"]):
        identity.cell(i, 0).text = label
        identity.cell(i, 1).text = "________________________________________"
        for r in identity.cell(i, 0).paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(NAVY)
        for c in identity.rows[i].cells:
            set_cell_margins(c, 70, 80, 70, 80)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    r = note.add_run("Public-use aggregate mortality data • Reproducible Python workflow")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)
    doc.add_page_break()


def add_report_map(doc):
    doc.add_heading("Report map", level=1)
    p = doc.add_paragraph("The document follows the required submission structure. Major sections are separated visually and use Word heading styles for navigation.")
    items = [
        "1. Title Page", "2. Introduction", "3. Dataset Overview", "4. Methodology",
        "5. Python Code Sections", "6. Results and Findings", "7. Visualizations (described in text)",
        "8. Challenges and How They Were Handled", "9. Discussion", "10. Conclusion", "11. References",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


def section(doc, number, title, intro=None):
    p = doc.add_heading(f"{number}. {title}", level=1)
    add_bottom_border(p, color="8FB6D1", size="8")
    if intro:
        doc.add_paragraph(intro)


def subheading(doc, title, level=2):
    return doc.add_heading(title, level=level)


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        doc.add_paragraph(item, style=style)


def callout(doc, title, text, fill="EAF2F8"):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_cell_shading(c, fill)
    set_cell_margins(c, 120, 160, 120, 160)
    p = c.paragraphs[0]
    r = p.add_run(title + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers: Sequence[str], rows: Iterable[Sequence], caption=None, font_size=8.5):
    rows = list(rows)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = True
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.text = str(h)
        set_cell_shading(c, NAVY)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(font_size)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = "" if val is None else str(val)
            if i % 2 == 1:
                set_cell_shading(cells[j], "F4F7F9")
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[j], 65, 75, 65, 75)
            for r in cells[j].paragraphs[0].runs:
                r.font.size = Pt(font_size)
    if caption:
        p = doc.add_paragraph(caption, style="Caption Custom")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t


def add_code(doc, label, code, explanations):
    doc.add_heading(label, level=3)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_cell_shading(c, "F3F5F7")
    set_cell_margins(c, 90, 110, 90, 110)
    p = c.paragraphs[0]
    p.style = doc.styles["Code Block"]
    p.add_run("```python\n" + code.strip() + "\n```")
    p = doc.add_paragraph()
    r = p.add_run("Code explanation and output interpretation")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    for line in explanations:
        q = doc.add_paragraph(line, style="List Bullet")
        q.paragraph_format.left_indent = Inches(0.28)
        q.paragraph_format.first_line_indent = Inches(-0.16)


def add_figure(doc, filename, caption, width=6.7):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption, style="Caption Custom")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def references(doc, include_tf=False):
    refs = [
        f"Centers for Disease Control and Prevention, National Center for Health Statistics. Provisional COVID-19 Deaths by Sex and Age. Dataset page: {SOURCE_PAGE}",
        f"CDC data download used for source verification: {SOURCE_CSV}",
        f"CDC WONDER provisional mortality access noted by the publisher as the successor source: {CDC_WONDER}",
        "McKinney, W. pandas: Python Data Analysis Library. https://pandas.pydata.org/",
        "Harris, C. R. et al. (2020). Array programming with NumPy. Nature, 585, 357–362. https://numpy.org/",
        "Waskom, M. L. (2021). seaborn: statistical data visualization. Journal of Open Source Software, 6(60), 3021. https://seaborn.pydata.org/",
        "Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830. https://scikit-learn.org/",
        "Hunter, J. D. (2007). Matplotlib: A 2D Graphics Environment. Computing in Science & Engineering, 9(3), 90–95. https://matplotlib.org/",
    ]
    if include_tf:
        refs.append("Abadi, M. et al. (2016). TensorFlow: A system for large-scale machine learning. https://www.tensorflow.org/")
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")


def save_doc(doc, filename):
    path = OUT / filename
    doc.core_properties.title = filename.replace("_", " ").replace(".docx", "")
    doc.core_properties.subject = "Data Science with Python internship report"
    doc.core_properties.author = "Internship Project"
    doc.core_properties.last_modified_by = ""
    doc.core_properties.revision = 1
    doc.core_properties.comments = ""
    doc.core_properties.keywords = "COVID-19, CDC, data science, Python"
    doc.save(path)
    return path


COMMON_IMPORTS = '''from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns

DATA_PATH = Path("Provisional COVID-19 Deaths by Sex and Age.csv")'''

CLEAN_LOAD = '''df = pd.read_csv(
    "cleaned_covid_deaths.csv",
    parse_dates=["data_as_of", "start_date", "end_date"],
    low_memory=False,
)

EXCLUSIVE_AGE_GROUPS = [
    "Under 1 year", "1-4 years", "5-14 years", "15-24 years",
    "25-34 years", "35-44 years", "45-54 years", "55-64 years",
    "65-74 years", "75-84 years", "85 years and over",
]'''


def build_week1():
    doc = Document()
    configure_document(doc, 1)
    add_title_page(doc, "Data Acquisition, Cleaning, and Preprocessing", 1, "A defensible cleaning workflow for provisional COVID-19 mortality counts")
    add_report_map(doc)

    section(doc, 2, "Introduction")
    para(doc, "This report documents how I acquired, inspected, cleaned, and prepared the CDC dataset Provisional COVID-19 Deaths by Sex and Age. The question is not simply whether blank cells exist; it is whether each blank, unusual value, and aggregate row can be interpreted correctly before the data are used for trends, clustering, or prediction.")
    para(doc, "I selected the dataset because it is a large public-health table from the National Center for Health Statistics, it mixes monthly, yearly, and cumulative records, and it contains privacy-driven suppression. Those features make it a realistic cleaning task: an apparently simple decision such as filling missing counts with zero would alter the signal and could understate mortality in smaller jurisdictions and age groups.")
    para(doc, "The practical problem answered here is: how can the 137,700-row source file be converted into an analysis-ready table without inventing death counts, double-counting overlapping age groups, deleting valid high-burden jurisdictions, or leaking a partial reporting period into later work?")
    callout(doc, "Scope decision.", "The source is frozen as of 27 September 2023 and covers deaths through 23 September 2023. The work treats it as a historical public-use file, not as a source of current 2026 COVID-19 surveillance.")

    section(doc, 3, "Dataset Overview")
    para(doc, f"The attached CSV is a {S['source_file_mb']:.2f} MB extract from the CDC/NCHS public dataset. It contains {S['raw_rows']:,} records and {S['raw_columns']} source columns. The observation grain is a combination of reporting period, jurisdiction, sex, and age group; it is not one row per person.")
    add_table(doc, ["Dimension", "Observed values", "Interpretation"], [
        ["Reporting group", "By Month (123,930); By Year (11,016); By Total (2,754)", "Three grains are stacked in one file."],
        ["Jurisdiction", "54", "50 states plus DC, Puerto Rico, New York City, and United States."],
        ["Sex", "3", "All Sexes, Female, Male."],
        ["Age group", "17", "Contains overlapping schemes such as 0–17, 18–29, 30–39 and 15–24, 25–34, 35–44."],
        ["Count measures", "6", "COVID-19, all-cause, pneumonia, overlap, influenza, and combined P/I/C counts."],
        ["Data coverage", "1 Jan 2020–23 Sep 2023", "September 2023 and the 2023 annual record are partial."],
    ], "Table 1. Structure of the attached source file.")
    para(doc, "All source columns arrived as text when loaded defensively. The six count fields contain comma-formatted numbers, and three showed mixed-type warnings under default inference. Year and Month also contain blanks that are structurally correct for cumulative and annual records. The Footnote field identifies privacy suppression when one or more cells fall between 1 and 9.")
    add_table(doc, ["Field", "Missing count", "Missing %", "Decision"], [
        ["COVID-19 Deaths", "39,430", "28.63%", "Retain as unknown; add suppression flag."],
        ["Total Deaths", "19,509", "14.17%", "Retain as unknown; do not use as a zero denominator."],
        ["Pneumonia Deaths", "44,864", "32.58%", "Retain as unknown; add suppression flag."],
        ["Pneumonia and COVID-19 Deaths", "36,884", "26.79%", "Retain as unknown; add suppression flag."],
        ["Influenza Deaths", "26,688", "19.38%", "Retain as unknown; add suppression flag."],
        ["Pneumonia, Influenza, or COVID-19 Deaths", "44,233", "32.12%", "Retain as unknown; add suppression flag."],
    ], "Table 2. Missingness in the six measure columns.")

    section(doc, 4, "Methodology", "The steps below follow the Week 1 brief in order. Each action includes the reason for doing it and the effect it has on later analysis.")
    subheading(doc, "Step 1 — Acquire from a reliable public source")
    para(doc, "I used the attached CSV and checked its title, column schema, update date, and publisher against the CDC Socrata record. This matters because similarly named CDC tables exist at weekly and monthly grains. The identifier 9bhg-hcku matches the 16-column file used here. A SHA-256 checksum would be recorded in a production run so that the exact input can be reproduced.")
    subheading(doc, "Step 2 — Inspect before converting")
    para(doc, "The first pass loaded every field as text. That is deliberate: automatic type inference can silently interpret comma-formatted counts as strings in one chunk and numbers in another. I profiled shape, unique categories, null counts, date examples, duplicates, and the relationship between Group, Year, and Month before changing values.")
    subheading(doc, "Step 3 — Standardize names and data types")
    para(doc, "Column names were converted to snake_case so the code is consistent and typo-resistant. Dates were parsed with the explicit month/day/year format. Commas were removed only from the six known count columns, after which conversion used a nullable numeric type. Year and Month use pandas nullable integers so their structural blanks remain valid.")
    subheading(doc, "Step 4 — Classify missingness instead of treating every blank alike")
    para(doc, "The count blanks are censored values, not confirmed zeros. Every observed missing count occurs on a footnoted row, and the footnote says counts from 1 to 9 have been suppressed. I therefore preserved NaN and created one Boolean suppression flag per measure. By contrast, Year is correctly blank for 2,754 By Total rows, and Month is correctly blank for all 11,016 By Year plus 2,754 By Total rows. Those structural blanks require no repair.")
    subheading(doc, "Step 5 — Check inconsistencies and erroneous entries")
    para(doc, "I tested duplicate rows, duplicate grain keys, negative counts, impossible date order, COVID-19 counts greater than all-cause deaths, overlap counts greater than either component, and combined P/I/C counts below their component maximum. Every one of these checks returned zero violations. This is stronger evidence than relying on a summary table because it tests the relationships the variables must satisfy.")
    subheading(doc, "Step 6 — Detect outliers at a coherent grain")
    para(doc, "An IQR rule was applied only to cumulative, All Sexes, All Ages jurisdiction rows, after excluding the United States and the overlapping New York City geography. The upper cutoff was 60,705.625 COVID-19 deaths and flagged California, Florida, and Texas. I kept them. Their large counts are plausible consequences of population size and pandemic burden, not typographical errors. Deleting them would remove the jurisdictions that contribute the most deaths and distort every later model.")
    subheading(doc, "Step 7 — Prevent aggregation errors")
    para(doc, "The file includes totals alongside components and two overlapping age-band systems. I created a canonical mutually exclusive age sequence for later analysis: Under 1, 1–4, 5–14, 15–24, 25–34, 35–44, 45–54, 55–64, 65–74, 75–84, and 85+. Summations must also filter one Group and one Sex level. This decision prevents double-counting without deleting useful alternative bands from the clean master table.")
    subheading(doc, "Step 8 — Mark partial periods and export with an audit trail")
    para(doc, "Rows for September 2023 monthly data and 2023 annual data end on 23 September, producing 5,508 partial-period rows. I added is_partial_period so later trend and model code can exclude them explicitly. The clean export keeps all 137,700 source rows, adds nine transparent fields, and does not overwrite the raw file.")

    section(doc, 5, "Python Code Sections")
    add_code(doc, "Code 1 — Load the attached CSV without premature type inference", COMMON_IMPORTS + '''

raw = pd.read_csv(DATA_PATH, dtype=str, low_memory=False)
print(raw.shape)
print(raw.columns.tolist())
print(raw[["Group", "Year", "Month", "State", "Sex", "Age Group"]].head())''', [
        "The local filename exactly matches the supplied dataset, so the code can run with the CSV beside the notebook or script.",
        "Loading as text avoids the mixed-type warning that appears in Pneumonia Deaths, Pneumonia and COVID-19 Deaths, and Influenza Deaths under automatic inference.",
        "The observed shape is (137700, 16), confirming that no row was lost during acquisition.",
    ])
    add_code(doc, "Code 2 — Profile grain, missingness, and duplicates", '''grain_key = ["Group", "Year", "Month", "State", "Sex", "Age Group"]

profile = {
    "rows": len(raw),
    "columns": raw.shape[1],
    "groups": raw["Group"].value_counts(dropna=False).to_dict(),
    "states": raw["State"].nunique(),
    "sexes": raw["Sex"].nunique(),
    "age_groups": raw["Age Group"].nunique(),
    "exact_duplicates": int(raw.duplicated().sum()),
    "grain_duplicates": int(raw.duplicated(grain_key).sum()),
}
missing = raw.isna().sum().sort_values(ascending=False)
print(profile)
print(missing)''', [
        "The grain key tests whether the combination that should identify a row is unique; both exact and key duplicate counts are zero.",
        "The value counts reveal three stacked grains, which is why filtering Group is mandatory before aggregation.",
        "The output also confirms 54 jurisdictions, three sex categories, and 17 age groups.",
    ])
    add_code(doc, "Code 3 — Parse dates, counts, and nullable calendar fields", '''rename = {
    "Data As Of": "data_as_of", "Start Date": "start_date",
    "End Date": "end_date", "Group": "group", "Year": "year",
    "Month": "month", "State": "state", "Sex": "sex",
    "Age Group": "age_group", "COVID-19 Deaths": "covid_deaths",
    "Total Deaths": "total_deaths", "Pneumonia Deaths": "pneumonia_deaths",
    "Pneumonia and COVID-19 Deaths": "pneumonia_covid_deaths",
    "Influenza Deaths": "influenza_deaths",
    "Pneumonia, Influenza, or COVID-19 Deaths": "pic_deaths",
    "Footnote": "footnote",
}
count_cols = ["covid_deaths", "total_deaths", "pneumonia_deaths",
              "pneumonia_covid_deaths", "influenza_deaths", "pic_deaths"]

df = raw.rename(columns=rename).copy()
for col in ["data_as_of", "start_date", "end_date"]:
    df[col] = pd.to_datetime(df[col], format="%m/%d/%Y", errors="raise")
for col in ["year", "month"]:
    df[col] = pd.to_numeric(df[col], errors="coerce").astype("Int64")
for col in count_cols:
    df[col] = pd.to_numeric(df[col].str.replace(",", regex=False),
                            errors="coerce").astype("Float64")''', [
        "Explicit date parsing turns malformed dates into visible errors rather than ambiguous day/month interpretations.",
        "Only the documented count fields have commas removed, protecting labels such as state or age group from accidental edits.",
        "Nullable dtypes preserve meaningful blanks while making arithmetic and validation possible.",
    ])
    add_code(doc, "Code 4 — Preserve suppressed counts and identify partial periods", '''for col in count_cols:
    df[f"{col}_suppressed"] = df[col].isna() & df["footnote"].notna()

df["has_suppression"] = df[
    [f"{col}_suppressed" for col in count_cols]
].any(axis=1)

expected_month_end = df["start_date"] + pd.offsets.MonthEnd(0)
expected_year_end = pd.to_datetime(
    df["year"].astype("string") + "-12-31",
    format="%Y-%m-%d", errors="coerce"
)
df["is_partial_period"] = (
    (df["group"].eq("By Month") & (df["end_date"] < expected_month_end)) |
    (df["group"].eq("By Year") & (df["end_date"] < expected_year_end))
)
print(df["has_suppression"].sum(), df["is_partial_period"].sum())''', [
        "The flags separate privacy suppression from true zero counts; a suppressed value remains unknown within the interval 1–9.",
        "The file contains 97,896 footnoted rows and 5,508 rows from partial month/year periods.",
        "Keeping the raw missing value plus its reason supports sensitivity analysis later without pretending the exact count is known.",
    ])
    add_code(doc, "Code 5 — Run domain and consistency assertions", '''assert not df.duplicated().any()
assert not df.duplicated(["group", "year", "month", "state", "sex", "age_group"]).any()
assert (df[count_cols].dropna() >= 0).all().all()
assert not (df["start_date"] > df["end_date"]).any()
assert not (df["end_date"] > df["data_as_of"]).any()
assert not (df["covid_deaths"] > df["total_deaths"]).any()
assert not (df["pneumonia_covid_deaths"] > df["covid_deaths"]).any()
assert not (df["pneumonia_covid_deaths"] > df["pneumonia_deaths"]).any()
component_max = df[["covid_deaths", "pneumonia_deaths",
                    "influenza_deaths"]].max(axis=1)
assert not (df["pic_deaths"] < component_max).any()''', [
        "Assertions turn data-quality expectations into executable tests; a future changed extract will stop at the exact failed rule.",
        "This extract has no negative counts, invalid dates, duplicate grain keys, or impossible component relationships.",
        "Passing these checks does not prove every number is correct, but it eliminates several important classes of internal inconsistency.",
    ])
    add_code(doc, "Code 6 — Detect but do not automatically delete large jurisdiction counts", '''state_total = df.query(
    "group == 'By Total' and sex == 'All Sexes' and age_group == 'All Ages'"
).loc[lambda x: ~x["state"].isin(["United States", "New York City"])]

q1, q3 = state_total["covid_deaths"].quantile([0.25, 0.75])
iqr = q3 - q1
upper = q3 + 1.5 * iqr
flags = state_total.loc[state_total["covid_deaths"] > upper,
                        ["state", "covid_deaths"]]
print(upper)
print(flags.sort_values("covid_deaths", ascending=False))''', [
        "The IQR test is performed at one consistent grain so national totals, age subgroups, and monthly rows cannot create artificial outliers.",
        "The cutoff is 60,705.625; California, Florida, and Texas are flagged.",
        "These records are retained because the values are plausible high-population burdens, showing why an outlier flag is not the same as an error label.",
    ])
    add_code(doc, "Code 7 — Derive a safe rate and write the clean artifact", '''df["covid_share_of_total"] = np.where(
    df["total_deaths"].fillna(0).gt(0),
    df["covid_deaths"] / df["total_deaths"],
    np.nan,
)

output_path = Path("cleaned_covid_deaths.csv")
df.to_csv(output_path, index=False)
print(df.shape, output_path.resolve())''', [
        "The derived share is created only when a real positive denominator is present; suppressed or zero totals yield NaN rather than infinity.",
        "The export has 137,700 rows and 25 columns, preserving the source data and adding transparent audit fields.",
        "The raw CSV remains unchanged, allowing every downstream result to be traced back to the original extract.",
    ])

    section(doc, 6, "Results and Findings")
    bullets(doc, [
        "The file is structurally complete at its intended grain: 137,700 unique keys and zero duplicate rows.",
        "Count-field missingness ranges from 14.17% for Total Deaths to 32.58% for Pneumonia Deaths. The pattern is explained by NCHS privacy suppression; it is not random missingness.",
        "All 39,804 rows without a footnote have complete count fields. This alignment is the main reason I rejected mean, median, and zero imputation for the clean master table.",
        "No negative count, impossible date, component-overlap violation, or combined-measure violation was found.",
        "California, Florida, and Texas exceed the state-level IQR cutoff, but their records are valid and were retained.",
        "September 2023 and the 2023 yearly records are incomplete through 23 September; 5,508 rows are flagged so later time comparisons can exclude them.",
        "The final master table keeps all source rows. Cleaning changed representation and added metadata; it did not manufacture or discard mortality events.",
    ])
    add_table(doc, ["Quality test", "Result", "Action"], [
        ["Exact duplicate rows", "0", "No removal needed"],
        ["Duplicate grain keys", "0", "Key accepted"],
        ["Negative count values", "0", "No correction needed"],
        ["COVID-19 > total deaths", "0", "No correction needed"],
        ["Overlap > COVID-19 or pneumonia", "0", "No correction needed"],
        ["Invalid date order", "0", "No correction needed"],
        ["State-level IQR flags", "3", "Keep and document"],
        ["Partial-period rows", "5,508", "Flag; exclude when a complete period is required"],
    ], "Table 3. Outcome of the executable data-quality checks.")

    section(doc, 7, "Visualizations (described in text)")
    subheading(doc, "Figure 1 — Missingness profile")
    add_figure(doc, "w1_missingness.png", "Figure 1. Percentage of values missing in each count measure. All missing count cells occur on privacy-footnoted rows.")
    para(doc, "Technically, this is a horizontal bar chart with count-field names on the y-axis and missing percentage on the x-axis. The analytical point is that suppression is uneven by measure: pneumonia and the combined P/I/C field lose roughly one-third of values, while total deaths lose about one-seventh. A complete-case analysis across all six fields would therefore discard a large and non-random share of smaller cells.")
    subheading(doc, "Figure 2 — Jurisdiction-level high-count flags")
    add_figure(doc, "w1_outlier_distribution.png", "Figure 2. Distribution of cumulative COVID-19 deaths across 52 non-overlapping jurisdictions, with the 1.5×IQR cutoff.")
    para(doc, "Technically, the histogram uses one cumulative All Sexes, All Ages record per state/DC/Puerto Rico and marks 60,705.625 with a dashed line. Analytically, the right tail is expected in count data because jurisdictions differ greatly in population. The chart supports keeping California, Florida, and Texas and using log transformations or scale-free features in later modeling instead of deleting the tail.")

    section(doc, 8, "Challenges and How They Were Handled")
    add_table(doc, ["Dataset-specific challenge", "Why it matters", "Resolution"], [
        ["Comma-formatted counts and mixed inferred types", "Arithmetic can fail or silently skip text values.", "Load as strings; strip commas only in documented count fields; cast to Float64."],
        ["Suppression presented as blanks", "Zero-fill would undercount small cells and bias small jurisdictions.", "Keep NaN and add measure-specific suppression flags."],
        ["Structural Year/Month blanks", "Dropping all incomplete rows would delete valid annual and cumulative records.", "Validate missingness against Group instead of blanket deletion."],
        ["Overlapping age categories", "Adding every age row would double-count deaths.", "Preserve source rows but define a mutually exclusive analytical age scheme."],
        ["Aggregates mixed with components", "United States, All Sexes, and All Ages totals can be counted with their subgroups.", "Require an explicit grain filter before every aggregation."],
        ["Large valid counts resemble statistical outliers", "Automatic trimming would erase real burden.", "Flag within a coherent grain, inspect context, and retain."],
        ["Final month/year is partial", "Trend and model targets would be artificially low.", "Add is_partial_period and exclude it in complete-period analyses."],
    ], "Table 4. Concrete cleaning problems and resolutions.")

    section(doc, 9, "Discussion")
    para(doc, "The main preprocessing decision was to preserve uncertainty rather than hide it. Suppressed values are interval-censored: they are known to lie from 1 to 9, but the exact count is unavailable. Keeping NaN means a later user must make an explicit modeling decision. For a sensitivity analysis, 1, 5, and 9 can serve as lower, midpoint, and upper scenarios; none should replace the official clean value silently.")
    para(doc, "Outlier handling has a similar practical consequence. Removing California, Florida, and Texas would improve the symmetry of a plot but damage the question the dataset can answer. For count prediction, a log1p target, robust loss, population denominator, or standardized profile is more defensible. For reporting totals, the original count must remain.")
    para(doc, "The approach is limited by the source. These are provisional aggregate counts by jurisdiction of occurrence, not individual records. The file supplies no population denominators, testing volume, vaccination, comorbidities, or exposure data. COVID-19 share of all deaths is descriptive and is not a population mortality rate. Because the publisher stopped updating this table after 27 September 2023, it cannot support current surveillance without a newer CDC source.")
    para(doc, "If cleaning were skipped or done badly, downstream effects would be concrete: zero-filled suppressions would make small cells look disease-free; overlapping age bands would inflate totals; September 2023 would appear as an abrupt decline; random row splitting could leak the same time pattern into training and testing; and removing high-count states would produce models that work only for small jurisdictions.")

    section(doc, 10, "Conclusion")
    para(doc, "I converted a 137,700-row, 16-column CDC extract into a 25-column analysis-ready master table while retaining every source record. The workflow parsed dates and counts, distinguished structural blanks from privacy suppression, marked 5,508 partial-period rows, validated domain relationships, documented overlapping categories, and retained three plausible high-count jurisdiction outliers.")
    para(doc, "The next step is exploratory analysis on fixed, non-overlapping slices of the cleaned data. Week 2 will use complete monthly periods, one sex level at a time, and a mutually exclusive age sequence so that every plotted number has a defensible denominator and grain.")

    section(doc, 11, "References")
    references(doc)
    return save_doc(doc, "Week_1_Data_Acquisition_Cleaning_Preprocessing.docx")


def build_week2():
    e = S["eda"]
    doc = Document(); configure_document(doc, 2)
    add_title_page(doc, "Exploratory Data Analysis and Visualization", 2, "Time, age, sex, respiratory overlap, and jurisdiction patterns")
    add_report_map(doc)
    section(doc, 2, "Introduction")
    para(doc, "This report uses the cleaned Week 1 version of the CDC Provisional COVID-19 Deaths by Sex and Age dataset to answer five specific questions: when national deaths peaked, how the cumulative burden varied by age, whether male and female counts followed the same age pattern, how COVID-19 moved with pneumonia and all-cause mortality, and how the COVID-19 share of total deaths differed across jurisdictions.")
    para(doc, "The dataset was chosen because the CDC provides the same mortality concepts at monthly, annual, and cumulative grains. That breadth supports useful comparisons, but only after the Week 1 controls for suppression, overlapping age groups, aggregate rows, and partial September 2023 data are carried forward.")
    callout(doc, "Analytical guardrail.", "Every visualization is built from one declared grain. Counts from All Ages are never added to age subgroups, All Sexes is never summed with Female and Male, and New York City is excluded from state comparisons because it overlaps New York.")
    section(doc, 3, "Dataset Overview")
    para(doc, "The raw file has 137,700 rows and 16 columns; the Week 1 clean artifact has 25 columns after parsed types, six suppression flags, a row-level suppression flag, a partial-period flag, and a COVID-19 share field are added. Data are current only through 23 September 2023. The September monthly and 2023 annual rows remain present but are marked partial.")
    add_table(doc, ["EDA slice", "Filter", "Rows / units", "Reason"], [
        ["National monthly trend", "By Month; United States; All Sexes; All Ages", "45 months; 44 complete", "One national series without subgroup double counting."],
        ["Age profile", "By Total; United States; All Sexes; 11 exclusive age bands", "11", "Age bands cover the life course once."],
        ["Sex-by-age", "By Total; United States; Female/Male; 11 exclusive age bands", "22", "Direct within-age comparison."],
        ["Jurisdiction share", "By Total; All Sexes; All Ages; exclude US and NYC", "52", "Non-overlapping state/DC/PR profiles."],
        ["Correlation", "44 complete national months", "44", "Avoid partial September 2023 and compare aligned series."],
    ], "Table 1. Analytical slices used in Week 2.")
    para(doc, "Count fields remain right-skewed, and suppression is concentrated in small cells. The chosen national and cumulative slices are complete for the principal measures. Correlations use Spearman's rank coefficient because the wave-shaped series are non-normal and include extreme pandemic peaks.")

    section(doc, 4, "Methodology", "The EDA follows the brief's key steps: choose a suitable public dataset, summarize it, transform only where needed, create multiple visualizations, and interpret each pattern rather than stopping at chart description.")
    subheading(doc, "Step 1 — Load the Week 1 clean artifact")
    para(doc, "I reused the parsed dates, suppression flags, and partial-period marker instead of recoding the source. This protects continuity: the EDA cannot accidentally reinterpret a censored count as zero or include the truncated September 2023 month in peak/low comparisons.")
    subheading(doc, "Step 2 — Fix one grain for each question")
    para(doc, "Each question has a filter contract shown in Table 1. That contract is part of the method, not just code convenience. The data contain multiple valid totals, so a chart can look polished while being numerically wrong if All Sexes or All Ages is combined with components.")
    subheading(doc, "Step 3 — Summarize central values and extremes")
    para(doc, "For national months I calculated death counts and COVID-19 as a percentage of total deaths. For age groups I calculated each band's share of all COVID-19 deaths and COVID-19 as a percentage of all deaths within the band. For jurisdictions I used COVID-19 divided by total deaths; I call it a death-share, not a population risk rate.")
    subheading(doc, "Step 4 — Transform and aggregate transparently")
    para(doc, "Age bands were ordered categorically rather than alphabetically. Male-to-female ratios were calculated within each age band. State shares were ranked only after the United States and New York City were removed. September 2023 stayed on the timeline for visual context but was shaded and excluded from complete-month minima and correlations.")
    subheading(doc, "Step 5 — Build annotated visualizations")
    para(doc, "I used a line chart for time, a count-plus-share dual-axis chart for age, grouped bars for sex, horizontal bars for jurisdiction ranking, and a heatmap for correlations. Titles, units, labels, legends, and partial-period annotations are included so the chart is understandable outside the surrounding text.")
    subheading(doc, "Step 6 — Interpret, cross-check, and avoid causal claims")
    para(doc, "The visual interpretation is cross-checked against exact grouped values. I describe associations and distributional differences, not causes. This matters because the aggregate file has no population denominators, vaccination, infection prevalence, or individual risk factors.")

    section(doc, 5, "Python Code Sections")
    add_code(doc, "Code 1 — Imports and continuity load", COMMON_IMPORTS + "\n\n" + CLEAN_LOAD, [
        "The code reads the Week 1 output, retaining its real snake_case field names and parsed date columns.",
        "The explicit age list is mutually exclusive; alternative overlapping bands such as 0–17, 18–29, and 30–39 are deliberately omitted from age sums.",
        "The resulting table still contains all 137,700 rows, while each analysis below selects a coherent subset.",
    ])
    add_code(doc, "Code 2 — Build the complete national monthly series", '''national_monthly = df.query(
    "group == 'By Month' and state == 'United States' "
    "and sex == 'All Sexes' and age_group == 'All Ages'"
).sort_values("start_date")

complete_months = national_monthly.loc[~national_monthly["is_partial_period"]].copy()
complete_months["covid_pct_total"] = (
    100 * complete_months["covid_deaths"] / complete_months["total_deaths"]
)
peak = complete_months.loc[complete_months["covid_deaths"].idxmax()]
print(peak[["start_date", "covid_deaths", "covid_pct_total"]])''', [
        "The filters produce one row per national month and prevent age/sex aggregate duplication.",
        "The partial-period flag excludes September 2023 from peak and correlation calculations while allowing it to be shown as incomplete context.",
        "January 2021 is the maximum complete month with 105,565 COVID-19 deaths, equal to 28.25% of all reported deaths that month.",
    ])
    add_code(doc, "Code 3 — Calculate the cumulative age profile", '''age = df.query(
    "group == 'By Total' and state == 'United States' and sex == 'All Sexes'"
).loc[lambda x: x["age_group"].isin(EXCLUSIVE_AGE_GROUPS)].copy()

age["age_group"] = pd.Categorical(
    age["age_group"], categories=EXCLUSIVE_AGE_GROUPS, ordered=True
)
age = age.sort_values("age_group")
age["share_of_covid_pct"] = 100 * age["covid_deaths"] / age["covid_deaths"].sum()
age["covid_pct_total"] = 100 * age["covid_deaths"] / age["total_deaths"]
print(age[["age_group", "covid_deaths", "share_of_covid_pct", "covid_pct_total"]])''', [
        "Categorical ordering places age groups in life-course order instead of lexical order.",
        "The 11 selected bands sum to the All Ages COVID-19 total, providing a direct aggregation check.",
        "The three 65+ bands account for 868,831 deaths, or 75.76% of the cumulative total.",
    ])
    add_code(doc, "Code 4 — Compare male and female counts within age", '''sex_age = df.query(
    "group == 'By Total' and state == 'United States' "
    "and sex in ['Female', 'Male']"
).loc[lambda x: x["age_group"].isin(EXCLUSIVE_AGE_GROUPS)]

sex_age = sex_age.pivot(index="age_group", columns="sex", values="covid_deaths")
sex_age = sex_age.reindex(EXCLUSIVE_AGE_GROUPS)
sex_age["male_to_female"] = sex_age["Male"] / sex_age["Female"]
print(sex_age.sort_values("male_to_female", ascending=False).head())''', [
        "Pivoting aligns Female and Male counts for the same age band, so each ratio has a meaningful numerator and denominator.",
        "Across the selected bands, male deaths total 629,728 versus 517,046 female deaths, a ratio of 1.218.",
        "The largest age-specific ratio is 1.75 in ages 45–54, indicating the biggest proportional sex gap occurs before the oldest ages.",
    ])
    add_code(doc, "Code 5 — Rank jurisdictions by COVID-19 share of all deaths", '''jurisdiction = df.query(
    "group == 'By Total' and sex == 'All Sexes' and age_group == 'All Ages'"
).loc[lambda x: ~x["state"].isin(["United States", "New York City"])].copy()

jurisdiction["covid_pct_total"] = (
    100 * jurisdiction["covid_deaths"] / jurisdiction["total_deaths"]
)
ranked = jurisdiction.sort_values("covid_pct_total", ascending=False)
print(ranked[["state", "covid_pct_total"]].head(15))''', [
        "This metric standardizes COVID-19 counts by all recorded deaths within each jurisdiction, reducing but not eliminating scale differences.",
        "New Jersey is highest at 11.40%, Vermont is lowest at 4.10%, and the 52-jurisdiction median is 8.95%.",
        "The result is a composition measure, not an age-standardized or population-based mortality rate, so it should not be labeled risk.",
    ])
    add_code(doc, "Code 6 — Measure rank correlations and draw the heatmap", '''measure_cols = [
    "covid_deaths", "total_deaths", "pneumonia_deaths",
    "pneumonia_covid_deaths", "influenza_deaths", "pic_deaths"
]
corr = complete_months[measure_cols].corr(method="spearman")

sns.heatmap(corr, annot=True, fmt=".2f", cmap="vlag", center=0,
            vmin=-1, vmax=1, square=True)
plt.title("Correlation among national monthly death-count series")
plt.tight_layout()
plt.show()''', [
        "Spearman correlation measures whether series rise and fall in a similar rank order without assuming a linear Gaussian relationship.",
        "COVID-19 and pneumonia deaths have a strong positive coefficient of 0.919; COVID-19 and total deaths are also strongly associated at 0.845.",
        "The coefficients describe synchronized monthly movement, not proof that one series causes the other.",
    ])

    section(doc, 6, "Results and Findings")
    add_table(doc, ["Finding", "Exact result", "Interpretation"], [
        ["Cumulative national COVID-19 deaths", "1,146,774", "9.32% of 12,303,399 total deaths in the file's cumulative period."],
        ["Peak complete month", "January 2021: 105,565", "COVID-19 represented 28.25% of all deaths that month."],
        ["August 2023", "3,844", "96.36% below the January 2021 peak; September is partial and not used for the comparison."],
        ["Age 65+ cumulative share", "868,831; 75.76%", "Three-quarters of deaths are concentrated in ages 65 and older."],
        ["Largest age count", "85+: 311,863", "The oldest band has the highest count."],
        ["Highest within-band COVID share", "65–74: 10.24%", "The largest count and largest within-band share are not the same concept."],
        ["Male/female total ratio", "1.218", "Male cumulative deaths exceed female deaths across the exclusive age scheme."],
        ["Largest sex ratio", "45–54: 1.75", "The proportional male gap peaks in midlife."],
        ["Jurisdiction range", "New Jersey 11.40%; Vermont 4.10%", "Death composition differs by 7.30 percentage points."],
        ["COVID/pneumonia monthly Spearman", "0.919", "The two respiratory mortality series move closely together over time."],
    ], "Table 2. Main EDA findings from reproducible filtered slices.")
    para(doc, "The time series is wave-shaped rather than a smooth decline. A single average would hide the January 2021 peak and later rebounds. The age results also separate two ideas: the 85+ group has the most COVID-19 deaths, but ages 65–74 have the highest COVID-19 share of that group's all-cause deaths. The grouped sex chart shows male excess across most adult ages, with the strongest proportional separation at 45–54.")
    para(doc, "Jurisdiction differences are material, but they should be treated as descriptive. The 11.40% New Jersey figure means COVID-19 was a larger component of all recorded deaths in this period than in Vermont; it does not adjust for age structure or population. The correlation heatmap indicates shared wave timing among COVID-19, pneumonia, and all-cause deaths and helps motivate lagged respiratory features in later forecasting.")

    section(doc, 7, "Visualizations (described in text)")
    add_figure(doc, "w2_national_monthly.png", "Figure 1. United States monthly COVID-19 deaths. September 2023 is shaded because it ends on 23 September.")
    para(doc, "Technical reading: a monthly line chart places Start Date on the x-axis and reported COVID-19 deaths on the y-axis; the highest complete point is annotated. Analytical reading: burden arrives in distinct waves, and January 2021 dominates the series. The shaded final month prevents a truncated reporting period from being mistaken for a confirmed decline.")
    add_figure(doc, "w2_age_profile.png", "Figure 2. Cumulative COVID-19 count (bars) and COVID-19 percentage of all deaths (line) by mutually exclusive age group.")
    para(doc, "Technical reading: bars use the left count axis, while the red line uses the right percentage axis. Analytical reading: counts rise sharply with age, but the share of all deaths peaks at 65–74 and then declines slightly in older bands because all-cause mortality also rises steeply. The two axes make that difference visible.")
    add_figure(doc, "w2_sex_age.png", "Figure 3. Male and female cumulative COVID-19 deaths within the same exclusive age groups.")
    para(doc, "Technical reading: paired bars compare sex within each ordered age band and use a shared death-count scale. Analytical reading: male deaths are generally higher, and the proportional gap is most pronounced at 45–54. Absolute female counts approach male counts at the oldest ages, which is consistent with the changing composition of the surviving population but cannot be causally resolved here.")
    add_figure(doc, "w2_state_share.png", "Figure 4. Fifteen highest jurisdictions by cumulative COVID-19 deaths as a percentage of total deaths.")
    para(doc, "Technical reading: a sorted horizontal bar chart shows the percentage on the x-axis and jurisdiction on the y-axis. Analytical reading: the top group is not simply the largest states, which shows why normalizing by total deaths changes the ranking. The chart still lacks age and population standardization, so it is a screening view rather than a performance league table.")
    add_figure(doc, "w2_correlation.png", "Figure 5. Spearman rank-correlation heatmap for six national monthly mortality series over complete months.")
    para(doc, "Technical reading: cells range from −1 to +1 and display coefficients. Analytical reading: COVID-19, pneumonia, pneumonia-plus-COVID, and the P/I/C union form a tightly moving respiratory block. Influenza is less synchronized, reflecting a different seasonal and pandemic-era pattern.")

    section(doc, 8, "Challenges and How They Were Handled")
    bullets(doc, [
        "Overlapping age bands made a naïve age sum invalid. I selected the 11-band mutually exclusive scheme and verified that it reconciles to All Ages.",
        "September 2023 ends on the 23rd. I shaded it in the trend but excluded it from peak, minimum, and correlation calculations.",
        "New York City overlaps the New York state geography. It was removed from cross-jurisdiction ranking while New York remained.",
        "Suppression is non-random in small cells. National and cumulative slices minimize the issue; where suppression remains, values are not silently converted to zero.",
        "Large count ranges could visually flatten small age groups. A companion percentage line and explicit table provide context instead of hiding the count scale.",
        "A high correlation could be misread as causation. The narrative uses association language and identifies shared timing and overlapping respiratory definitions.",
    ])

    section(doc, 9, "Discussion")
    para(doc, "Practically, the results point to an age-concentrated but not age-exclusive burden. The 65+ population accounts for 75.76% of deaths, so long-term-care readiness and older-adult protection remain central when studying the historical period. At the same time, the 1.75 male-to-female ratio at 45–54 identifies a subgroup whose proportional gap would disappear if analysis focused only on the oldest-age counts.")
    para(doc, "The monthly wave pattern means a model should be evaluated in chronological blocks. A random train/test split would distribute the same pandemic wave across both sets and overstate generalization. The strong COVID-pneumonia association suggests lagged pneumonia may carry predictive information, but overlapping definitions and concurrent coding mean it must not be interpreted as an independent causal driver.")
    para(doc, "Limitations are substantial. The data are aggregate, provisional, and frozen in September 2023. They record occurrence rather than individual exposure, contain interval-censored small counts, and provide no denominators for population risk or age standardization. Visual differences cannot identify policy effects, infection rates, or biological mechanisms.")
    para(doc, "If EDA were skipped, later clustering could be driven by raw population size, a model could train on incomplete September, and overlapping age rows could create impossible targets. The transformations used here—grain filters, exclusive age categories, complete-period flags, and scale-aware shares—directly define the safe feature set for Weeks 3 and 4.")

    section(doc, 10, "Conclusion")
    para(doc, "The EDA found 1,146,774 cumulative U.S. COVID-19 deaths, a January 2021 monthly peak of 105,565, a 75.76% share among ages 65+, a cumulative male/female ratio of 1.218, and wide jurisdiction variation in COVID-19's share of all deaths. COVID-19 and pneumonia monthly deaths have a Spearman correlation of 0.919, reinforcing the respiratory-wave pattern.")
    para(doc, "The next stage will cluster 52 non-overlapping jurisdictions using standardized, scale-free mortality profiles. The cluster features will use age share, sex share, respiratory overlap, and wave timing rather than raw death counts so population size does not decide the groups by itself.")
    section(doc, 11, "References"); references(doc)
    return save_doc(doc, "Week_2_Exploratory_Data_Analysis_Visualization.docx")


def build_week3():
    c = S["clustering"]
    doc = Document(); configure_document(doc, 3)
    add_title_page(doc, "Unsupervised Learning and Clustering Analysis", 3, "Segmenting 52 jurisdiction mortality profiles with K-Means")
    add_report_map(doc)
    section(doc, 2, "Introduction")
    para(doc, "This report groups U.S. states, the District of Columbia, and Puerto Rico according to how COVID-19 mortality is distributed across all deaths, pneumonia overlap, influenza, older age, sex, and pandemic year. The aim is descriptive segmentation: identify jurisdictions with similar profiles even when their raw population and death counts differ.")
    para(doc, "The cleaned CDC dataset is suitable because it supplies the same cumulative and yearly measures for every jurisdiction. Week 2 showed that raw counts are heavily affected by jurisdiction size, so the clustering question is framed around seven percentages. The practical question is whether a small number of recurring mortality profiles can guide targeted comparison, monitoring, and hypothesis generation.")
    callout(doc, "Interpretation boundary.", f"The selected four-cluster solution has a silhouette score of {c['best_silhouette']:.3f}, which is modest. The clusters are useful descriptive neighborhoods, not natural laws or causal categories.")
    section(doc, 3, "Dataset Overview")
    para(doc, "The source file contains 137,700 rows and 16 raw columns; analysis starts from the 25-column Week 1 clean artifact. I excluded the national total and New York City. The United States is an aggregate, while New York City overlaps New York; keeping either in the jurisdiction matrix would distort distance and independence. The final modeling unit is 52 jurisdictions.")
    add_table(doc, ["Feature", "Definition", "Why included"], [
        ["COVID % of all deaths", "100 × COVID-19 / total deaths", "Overall COVID composition without raw size."],
        ["Pneumonia % of all deaths", "100 × pneumonia / total deaths", "Respiratory mortality context."],
        ["COVID with pneumonia %", "100 × pneumonia+COVID / COVID", "Clinical/coding overlap profile."],
        ["Influenza % of all deaths", "100 × influenza / total deaths", "Separates non-COVID respiratory burden."],
        ["Age 65+ COVID %", "65–74 + 75–84 + 85+ COVID / all-age COVID", "Age concentration."],
        ["Male COVID %", "Male / (Male + Female) COVID", "Sex composition."],
        ["2021 % of 2020–22 COVID", "2021 / sum of 2020, 2021, 2022", "Wave timing without partial 2023."],
    ], "Table 1. Seven scale-free features used for clustering.")
    para(doc, "The cumulative fields used for features are observed for all 52 units, so no suppressed model feature required imputation. Raw features use different percentage ranges and variances; StandardScaler centers each at zero and scales to unit variance before Euclidean distance is calculated.")

    section(doc, 4, "Methodology", "The workflow follows the required order: select an appropriate unit, preprocess to improve clustering, apply and tune K-Means, visualize the solution, and translate cluster profiles into practical meaning.")
    subheading(doc, "Step 1 — Define an independent analytical unit")
    para(doc, "Each row of the feature matrix represents one non-overlapping jurisdiction. Excluding the national total prevents a single aggregate from sitting far from every state, and excluding New York City prevents one region from being represented twice.")
    subheading(doc, "Step 2 — Engineer interpretable percentage features")
    para(doc, "I avoided raw counts because they would mostly separate California and Texas from smaller states. Each feature answers a distinct composition or timing question. The age numerator uses exactly three mutually exclusive older bands, and the year-timing feature stops at 2022 because 2023 is incomplete.")
    subheading(doc, "Step 3 — Check completeness, correlation, and scale")
    para(doc, "All 364 feature cells are observed. The seven variables are standardized because a one-percentage-point difference has different meaning across low-variance influenza share and higher-variance pneumonia overlap. Standardization gives each feature an opportunity to influence distance while retaining its direction.")
    subheading(doc, "Step 4 — Compare candidate cluster counts")
    para(doc, "K-Means was run for k=2 through k=8 with 50 random initializations per candidate and random_state=42. I examined inertia for compactness and silhouette for separation. Silhouette reaches its highest value at k=4 (0.222); k=7 is close at 0.220 but creates more, smaller groups with less stable practical interpretation. Four clusters are therefore selected.")
    subheading(doc, "Step 5 — Fit the final model and visualize structure")
    para(doc, "The final four-cluster K-Means uses 100 starts to reduce sensitivity to poor centroid initialization. PCA projects the standardized seven-dimensional matrix to two dimensions for display; the first two components retain 62.19% of variance. A Ward-linkage dendrogram provides a second view of similarity without changing the K-Means assignments.")
    subheading(doc, "Step 6 — Name clusters from centroid profiles")
    para(doc, "Labels are based on cluster means relative to the full 52-jurisdiction mean, not on geography alone. I report membership and exact centroid values, then give plain-language names. These names summarize observed profiles and should not be used to rank policy performance.")

    section(doc, 5, "Python Code Sections")
    add_code(doc, "Code 1 — Imports and clean data load", '''from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from scipy.cluster.hierarchy import dendrogram, linkage
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from sklearn.metrics import silhouette_score
from sklearn.preprocessing import StandardScaler

df = pd.read_csv("cleaned_covid_deaths.csv",
                 parse_dates=["data_as_of", "start_date", "end_date"])
jurisdictions = sorted(set(df["state"]) - {"United States", "New York City"})''', [
        "The code continues from the Week 1 artifact and uses the actual clean field names.",
        "Removing United States and New York City leaves 52 non-overlapping modeling units: 50 states, DC, and Puerto Rico.",
        "All subsequent joins use jurisdiction name as the index, making row alignment explicit.",
    ])
    add_code(doc, "Code 2 — Construct the seven-feature jurisdiction matrix", '''base = df.query(
    "group == 'By Total' and sex == 'All Sexes' and age_group == 'All Ages'"
).loc[lambda x: x["state"].isin(jurisdictions)].set_index("state")

features = pd.DataFrame(index=jurisdictions)
features["covid_pct_all_deaths"] = 100 * base["covid_deaths"] / base["total_deaths"]
features["pneumonia_pct_all_deaths"] = 100 * base["pneumonia_deaths"] / base["total_deaths"]
features["covid_with_pneumonia_pct"] = 100 * base["pneumonia_covid_deaths"] / base["covid_deaths"]
features["influenza_pct_all_deaths"] = 100 * base["influenza_deaths"] / base["total_deaths"]

age65 = df.query(
    "group == 'By Total' and sex == 'All Sexes' "
    "and age_group in ['65-74 years', '75-84 years', '85 years and over']"
).loc[lambda x: x["state"].isin(jurisdictions)].groupby("state")["covid_deaths"].sum()
features["age_65plus_covid_pct"] = 100 * age65 / base["covid_deaths"]''', [
        "The first four features are cumulative burden/overlap percentages, and the fifth measures older-age concentration.",
        "All divisions use like-for-like cumulative records at one sex and age grain.",
        "Using shares prevents high-population states from dominating solely because their counts are larger.",
    ])
    add_code(doc, "Code 3 — Add sex and wave-timing composition", '''sex = df.query(
    "group == 'By Total' and age_group == 'All Ages' and sex in ['Female', 'Male']"
).loc[lambda x: x["state"].isin(jurisdictions)]
sex = sex.pivot(index="state", columns="sex", values="covid_deaths")
features["male_covid_pct"] = 100 * sex["Male"] / (sex["Male"] + sex["Female"])

yearly = df.query(
    "group == 'By Year' and sex == 'All Sexes' and age_group == 'All Ages' "
    "and year in [2020, 2021, 2022]"
).loc[lambda x: x["state"].isin(jurisdictions)]
yearly = yearly.pivot(index="state", columns="year", values="covid_deaths")
features["covid_2021_pct_2020_22"] = 100 * yearly[2021] / yearly.sum(axis=1)

assert features.shape == (52, 7)
assert features.notna().all().all()''', [
        "The sex feature measures composition rather than the male/female count ratio, keeping it on a bounded percentage scale.",
        "The wave feature excludes partial 2023 and records how concentrated 2020–22 burden was in 2021.",
        "The assertions confirm the intended 52 × 7 matrix with no imputed feature cells.",
    ])
    add_code(doc, "Code 4 — Standardize and select k with silhouette plus inertia", '''X = StandardScaler().fit_transform(features)
selection = []
for k in range(2, 9):
    model = KMeans(n_clusters=k, n_init=50, random_state=42)
    labels = model.fit_predict(X)
    selection.append({
        "k": k,
        "inertia": model.inertia_,
        "silhouette": silhouette_score(X, labels),
    })

selection = pd.DataFrame(selection)
best_k = int(selection.loc[selection["silhouette"].idxmax(), "k"])
print(selection)
print("Selected k:", best_k)''', [
        "Testing multiple k values avoids choosing the number of groups by appearance alone.",
        "Inertia falls as k increases by construction; silhouette adds a separation criterion that can peak.",
        "The maximum silhouette is 0.222 at k=4, compared with 0.220 at k=7, so the simpler four-cluster solution is preferred.",
    ])
    add_code(doc, "Code 5 — Fit final K-Means, profile clusters, and project with PCA", '''kmeans = KMeans(n_clusters=4, n_init=100, random_state=42)
features["cluster"] = kmeans.fit_predict(X)

# Stable report labels: order cluster IDs by mean COVID share.
order = features.groupby("cluster")["covid_pct_all_deaths"].mean().sort_values().index
remap = {old: new for new, old in enumerate(order)}
features["cluster"] = features["cluster"].map(remap)

profile = features.groupby("cluster").mean(numeric_only=True)
pca = PCA(n_components=2)
features[["pc1", "pc2"]] = pca.fit_transform(X)
print(profile)
print("PCA variance retained:", pca.explained_variance_ratio_.sum())''', [
        "Using 100 initializations reduces dependence on a single random centroid start.",
        "Cluster IDs are arbitrary, so remapping by mean COVID share makes the report order reproducible without changing membership.",
        "The first two principal components retain 62.19% of standardized feature variance; the projection is informative but not a complete representation.",
    ])
    add_code(doc, "Code 6 — Create a hierarchical robustness view", '''Z = linkage(X, method="ward")
plt.figure(figsize=(12, 5))
dendrogram(Z, labels=features.index, leaf_rotation=90, leaf_font_size=7)
plt.ylabel("Ward linkage distance")
plt.title("Hierarchical view of jurisdiction profile similarity")
plt.tight_layout()
plt.show()''', [
        "Ward linkage merges groups that add the least within-cluster variance, providing a non-K-Means view of the same standardized matrix.",
        "The dendrogram does not assign the final labels; it is used to inspect whether similar neighborhoods persist under another method.",
        "Several branches align with the four-profile interpretation, while close merge heights reinforce that boundaries are not sharp.",
    ])

    section(doc, 6, "Results and Findings")
    add_table(doc, ["Cluster", "n", "Plain-language profile", "Members"], [
        ["0", "8", "Lower COVID share; male-skewed; 2021-concentrated", ", ".join(c["cluster_members"]["0"])],
        ["1", "13", "Older-age concentration; lower pneumonia overlap; less 2021-concentrated", ", ".join(c["cluster_members"]["1"])],
        ["2", "23", "Broad higher-burden middle profile", ", ".join(c["cluster_members"]["2"])],
        ["3", "8", "High pneumonia/overlap; younger and male-skewed profile", ", ".join(c["cluster_members"]["3"])],
    ], "Table 2. Final four clusters and jurisdiction membership.", font_size=7.2)
    add_table(doc, ["Cluster", "COVID % all", "Pneumonia % all", "COVID with pneumonia %", "65+ COVID %", "Male COVID %", "2021 % of 2020–22"], [
        ["0", "7.125", "7.680", "50.409", "74.662", "57.669", "47.838"],
        ["1", "8.130", "8.040", "39.041", "82.672", "53.565", "35.012"],
        ["2", "9.265", "9.260", "47.463", "76.435", "53.279", "42.591"],
        ["3", "9.355", "11.534", "63.144", "70.803", "57.434", "43.162"],
    ], "Table 3. Selected centroid values; influenza share is included in modeling but omitted here for space.")
    para(doc, "Cluster 0 contains Alaska, several Mountain/Western states, Hawaii, Oregon, and Washington. Its COVID share is lowest at 7.125%, yet 47.838% of 2020–22 COVID deaths occurred in 2021 and male share is high at 57.669%. The practical profile is lower cumulative composition but later/concentrated and male-skewed burden.")
    para(doc, "Cluster 1 has the oldest profile: 82.672% of COVID deaths are among ages 65+, pneumonia overlap is the lowest at 39.041%, and only 35.012% of 2020–22 deaths fall in 2021. This is an older-age and less-2021-concentrated profile containing many New England and Upper Midwest jurisdictions.")
    para(doc, "Cluster 2 is the largest group (23). Its COVID and pneumonia shares are both about 9.26%, older-age concentration is 76.435%, and sex composition is less male-skewed than Clusters 0 and 3. It represents a broad higher-burden middle pattern rather than an extreme on one feature.")
    para(doc, "Cluster 3 combines the highest pneumonia share (11.534%), highest COVID-with-pneumonia overlap (63.144%), lowest 65+ share (70.803%), and high male share (57.434%). Arizona, California, DC, Florida, Nevada, New Mexico, Puerto Rico, and Texas form this younger, male-skewed, respiratory-overlap profile.")

    section(doc, 7, "Visualizations (described in text)")
    add_figure(doc, "w3_k_selection.png", "Figure 1. Inertia and silhouette for k=2 through k=8; k=4 has the highest silhouette.")
    para(doc, "Technical reading: the left axis shows within-cluster sum of squares and the right axis shows silhouette, both by k. Analytical reading: compactness improves continuously, but separation does not. Four groups offer the best measured separation with a much simpler interpretation than the near-tied seven-group alternative.")
    add_figure(doc, "w3_pca_clusters.png", "Figure 2. Four K-Means groups in two-dimensional PCA space; the display retains 62.19% of variance.")
    para(doc, "Technical reading: each point is a standardized jurisdiction profile, color indicates cluster, and axes are the first two principal components. Analytical reading: groups occupy recognizable regions but overlap at boundaries, which agrees with the modest silhouette and argues against treating membership as absolute.")
    add_figure(doc, "w3_cluster_heatmap.png", "Figure 3. Cluster means expressed as standard deviations above or below the 52-jurisdiction mean.")
    para(doc, "Technical reading: columns are the seven features, rows are clusters, and color/annotation give standardized centroid deviations. Analytical reading: Cluster 1 is defined most strongly by older-age concentration and low overlap, while Cluster 3 is defined by high pneumonia and overlap plus a younger/male-skewed profile. This reveals the feature combination behind each label.")
    add_figure(doc, "w3_dendrogram.png", "Figure 4. Ward-linkage dendrogram for the same standardized feature matrix.")
    para(doc, "Technical reading: lower merge height means greater profile similarity. Analytical reading: several tight jurisdiction pairs and subgroups exist, but branch heights rise gradually rather than showing one overwhelming cut. The data support useful segmentation, not a single indisputable taxonomy.")

    section(doc, 8, "Challenges and How They Were Handled")
    bullets(doc, [
        "Raw counts were dominated by population size. I replaced them with seven interpretable percentages before standardization.",
        "New York City overlaps New York and the national total aggregates every jurisdiction. Both were excluded to preserve independent units.",
        "The candidate silhouette scores were low and close. I reported them, chose the simplest maximum at k=4, and avoided claiming strong natural separation.",
        "Cluster numbers changed meaning across random runs. I used n_init=100, random_state=42, and remapped labels by mean COVID share for stable reporting.",
        "PCA loses 37.81% of variance. It is used only for visualization; K-Means is fit in all seven standardized dimensions.",
        "2023 is incomplete. Wave timing uses only 2020–2022 so a partial year cannot reduce the denominator inconsistently.",
    ])

    section(doc, 9, "Discussion")
    para(doc, "The segmentation can support comparison sets. Cluster 1 jurisdictions should prioritize older-adult and long-term-care comparisons because their age concentration is the defining feature. Cluster 3 comparisons should examine pneumonia coding/clinical overlap and working-age male patterns. Cluster 0's 2021 concentration suggests that calendar timing matters when evaluating interventions, while Cluster 2 provides a large reference group for broad burden patterns.")
    para(doc, "These are research and monitoring implications, not causal diagnoses. Percentage features can still reflect age structure, coding practice, care access, and epidemic timing. A jurisdiction can move clusters if revised provisional counts change a feature. The model contains 52 observations for seven variables, so small cluster differences should be validated with bootstrapping or newer data before operational use.")
    para(doc, "If preprocessing were skipped, raw totals would mostly create a population-size clustering. If standardization were skipped, high-variance pneumonia overlap could dominate influenza share. If k were chosen only from a visually appealing PCA plot, the decision would ignore the 37.81% of variance outside the display. Each safeguard directly protects interpretability.")

    section(doc, 10, "Conclusion")
    para(doc, "K-Means on seven standardized mortality-profile percentages produced four interpretable groups of sizes 8, 13, 23, and 8. The selected k has the highest tested silhouette (0.222), and the first two principal components retain 62.19% of variance. Cluster profiles differ most clearly in older-age concentration, male share, pneumonia overlap, overall COVID composition, and concentration in 2021.")
    para(doc, "The next step is supervised one-month-ahead forecasting. Cluster labels will not be used as targets; instead, chronological lag features and jurisdiction identity will predict monthly COVID-19 counts, with performance compared against a previous-month baseline.")
    section(doc, 11, "References"); references(doc)
    return save_doc(doc, "Week_3_Unsupervised_Learning_Clustering.docx")


def build_week4():
    m = S["supervised"]
    doc = Document(); configure_document(doc, 4)
    add_title_page(doc, "Supervised Learning Model Implementation", 4, "One-month-ahead jurisdiction COVID-19 death forecasting")
    add_report_map(doc)
    section(doc, 2, "Introduction")
    para(doc, "This report frames the CDC data as a supervised regression problem: predict each jurisdiction's COVID-19 deaths one month ahead using only information available before the forecast month. The problem is useful for staffing and surveillance planning, but the exercise is also designed to expose where a conventional tree model fails when the pandemic regime changes.")
    para(doc, "The cleaned dataset was retained because it supplies monthly COVID-19, all-cause, and pneumonia counts for 52 non-overlapping jurisdictions from 2020 through August 2023. The key question is not whether a model can fit historical waves; it is whether it improves on the simple and operationally strong rule of using the previous month's count.")
    callout(doc, "Bottom line.", "The gradient-boosted model is not deployment-ready. On the January–August 2023 test period it underperforms the previous-month baseline on MAE, RMSE, R², and WAPE. Reporting that failure is part of correct model evaluation, not a reason to hide the holdout result.", fill="FBE9E7")
    section(doc, 3, "Dataset Overview")
    para(doc, "The 137,700-row source mixes monthly, yearly, and cumulative records. Modeling uses only By Month, All Sexes, All Ages rows for 52 jurisdictions; United States and New York City are excluded for the same independence reasons used in Week 3. Partial September 2023 is excluded.")
    add_table(doc, ["Stage", "Rows", "Period", "Purpose"], [
        ["Feature-ready panel", "2,069", "Apr 2020–Aug 2023", "Rows with three lags, rolling history, observed target, and lagged covariates."],
        ["Final training", "1,690", "Apr 2020–Dec 2022", "Fit model before the holdout period."],
        ["Final test", "379", "Jan–Aug 2023", "Strict future evaluation; suppressed targets are excluded."],
        ["CV fold 1", "760 train / 310 validation", "Validate Jul–Dec 2021", "Expansion-window check during a high-wave regime."],
        ["CV fold 2", "1,070 / 311", "Validate Jan–Jun 2022", "Check transition from winter peak."],
        ["CV fold 3", "1,381 / 309", "Validate Jul–Dec 2022", "Check lower, smoother late-2022 regime."],
    ], "Table 1. Chronological split design.")
    para(doc, "For lag predictors only, a suppressed count is represented by the midpoint 5 of the known interval 1–9. A row with a suppressed target is never scored as though 5 were observed; it is removed from training/evaluation. This separates a bounded predictor approximation from ground-truth fabrication.")

    section(doc, 4, "Methodology", "The workflow defines the supervised problem, preprocesses and engineers leak-free features, trains a selected model, uses expanding-time validation, and evaluates several metrics plus failure cases.")
    subheading(doc, "Step 1 — Define the target and forecast horizon")
    para(doc, "The target is the current row's monthly COVID-19 death count, while every count predictor is shifted by at least one month within jurisdiction. The task is therefore a one-step-ahead forecast, not a contemporaneous explanation. September 2023 is removed because it covers only 23 days.")
    subheading(doc, "Step 2 — Engineer features without future information")
    para(doc, "Features are COVID-19 lags 1–3, the previous three-month mean, prior-month total deaths, prior-month pneumonia deaths, sine and cosine of month, a linear time index, and one-hot jurisdiction identity. Lagged respiratory and all-cause measures capture recent scale; cyclical terms represent season; one-hot identity allows persistent level differences without imposing an artificial state order.")
    subheading(doc, "Step 3 — Use chronological validation")
    para(doc, "Three expanding-window folds respect calendar order. Random cross-validation was rejected because adjacent months and pandemic waves are autocorrelated; random splitting would let future-regime examples inform earlier predictions. The latest complete months, January–August 2023, remain untouched for final testing.")
    subheading(doc, "Step 4 — Choose the model and transformations")
    para(doc, "I used HistGradientBoostingRegressor with shallow leaf structure, learning rate 0.05, 350 boosting iterations, minimum 12 samples per leaf, and L2 regularization. Gradient boosting captures nonlinear interactions among lags, season, and jurisdiction. The target is log1p-transformed to reduce domination by California/Texas and peak waves, then transformed back for count-scale metrics.")
    subheading(doc, "Step 5 — Benchmark against a hard-to-beat baseline")
    para(doc, "The baseline predicts that next month equals the previous month. In a smooth low-burden period this rule can be excellent. A useful model must outperform it, not merely show positive R². The same target rows are used for model and baseline metrics.")
    subheading(doc, "Step 6 — Evaluate trade-offs and inspect failures")
    para(doc, "MAE describes the typical absolute miss in deaths, RMSE penalizes large misses more strongly, R² measures variance explained relative to a mean predictor, and WAPE expresses total absolute error relative to total observed deaths. Error tables identify months and jurisdictions where summary metrics hide operationally serious misses.")

    section(doc, 5, "Python Code Sections")
    add_code(doc, "Code 1 — Imports, cleaned panel, and non-overlapping jurisdictions", '''import numpy as np
import pandas as pd
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import HistGradientBoostingRegressor
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler

df = pd.read_csv("cleaned_covid_deaths.csv",
                 parse_dates=["data_as_of", "start_date", "end_date"])
jurisdictions = sorted(set(df["state"]) - {"United States", "New York City"})
panel = df.query(
    "group == 'By Month' and sex == 'All Sexes' and age_group == 'All Ages'"
).loc[lambda x: x["state"].isin(jurisdictions) & ~x["is_partial_period"]]
panel = panel.sort_values(["state", "start_date"])''', [
        "The model uses one monthly all-age, all-sex series per non-overlapping jurisdiction.",
        "Sorting within state is required before lag construction; otherwise shift would connect unrelated jurisdictions.",
        "Removing partial periods prevents a 23-day September target from being learned as a full-month decline.",
    ])
    add_code(doc, "Code 2 — Create lagged and seasonal features", '''for col in ["covid_deaths", "total_deaths", "pneumonia_deaths"]:
    panel[f"{col}_for_lag"] = panel[col].fillna(5.0)

g = panel.groupby("state", observed=True)
panel["lag1_covid"] = g["covid_deaths_for_lag"].shift(1)
panel["lag2_covid"] = g["covid_deaths_for_lag"].shift(2)
panel["lag3_covid"] = g["covid_deaths_for_lag"].shift(3)
panel["rolling3_covid"] = g["covid_deaths_for_lag"].transform(
    lambda s: s.shift(1).rolling(3).mean()
)
panel["lag1_total"] = g["total_deaths_for_lag"].shift(1)
panel["lag1_pneumonia"] = g["pneumonia_deaths_for_lag"].shift(1)
panel["month_sin"] = np.sin(2 * np.pi * panel["month"] / 12)
panel["month_cos"] = np.cos(2 * np.pi * panel["month"] / 12)
panel["time_index"] = (panel["start_date"].dt.year - 2020) * 12 + panel["month"] - 1''', [
        "The midpoint 5 is used only when a lagged predictor is interval-suppressed; observed targets remain required for scoring.",
        "All rolling calculations shift first, so the current target cannot enter its own predictors.",
        "Sine and cosine encode December–January adjacency, unlike a raw month integer that places them far apart.",
    ])
    add_code(doc, "Code 3 — Define preprocessing and the gradient-boosted regressor", '''num_features = [
    "lag1_covid", "lag2_covid", "lag3_covid", "rolling3_covid",
    "lag1_total", "lag1_pneumonia", "month_sin", "month_cos", "time_index"
]
cat_features = ["state"]
model_data = panel.dropna(subset=num_features + ["covid_deaths"]).copy()

preprocess = ColumnTransformer([
    ("numeric", StandardScaler(), num_features),
    ("state", OneHotEncoder(handle_unknown="ignore", sparse_output=False), cat_features),
])
model = HistGradientBoostingRegressor(
    learning_rate=0.05, max_iter=350, max_leaf_nodes=15,
    min_samples_leaf=12, l2_regularization=1.0, random_state=42,
)
pipeline = Pipeline([("preprocess", preprocess), ("model", model)])''', [
        "One-hot encoding represents jurisdiction identity without treating state names as an ordered numeric variable.",
        "Regularization and minimum leaf size limit overly specific rules in a panel of only 2,069 usable observations.",
        "After encoding, the training design contains 61 features: nine numeric predictors and 52 jurisdiction indicators.",
    ])
    add_code(doc, "Code 4 — Run expanding-window validation", '''folds = [
    ("2021 H2", "2021-07-01", "2021-12-01"),
    ("2022 H1", "2022-01-01", "2022-06-01"),
    ("2022 H2", "2022-07-01", "2022-12-01"),
]
for label, start, end in folds:
    train = model_data[model_data["start_date"] < start]
    valid = model_data[model_data["start_date"].between(start, end)]
    pipeline.fit(train[num_features + cat_features], np.log1p(train["covid_deaths"]))
    pred = np.clip(np.expm1(pipeline.predict(valid[num_features + cat_features])), 0, None)
    baseline = valid["lag1_covid"].to_numpy()
    print(label, mean_absolute_error(valid["covid_deaths"], pred),
          mean_absolute_error(valid["covid_deaths"], baseline))''', [
        "Each fold trains only on dates earlier than its validation block, reproducing how the model would have been used at the time.",
        "The model beats the baseline in 2021 H2 and 2022 H1 but loses badly in 2022 H2, exposing regime-dependent performance.",
        "That instability is visible before the final 2023 test and argues against selecting a model from one favorable split.",
    ])
    add_code(doc, "Code 5 — Fit through 2022 and evaluate January–August 2023", '''train = model_data[model_data["start_date"] < "2023-01-01"]
test = model_data[model_data["start_date"].between("2023-01-01", "2023-08-01")]
X_cols = num_features + cat_features
pipeline.fit(train[X_cols], np.log1p(train["covid_deaths"]))
pred = np.clip(np.expm1(pipeline.predict(test[X_cols])), 0, None)
actual = test["covid_deaths"].to_numpy()

metrics = {
    "MAE": mean_absolute_error(actual, pred),
    "RMSE": mean_squared_error(actual, pred) ** 0.5,
    "R2": r2_score(actual, pred),
    "WAPE_pct": 100 * np.abs(actual - pred).sum() / actual.sum(),
}
print(metrics)''', [
        "The final fit uses all observed rows through December 2022 and never sees 2023 targets during training.",
        "Test results are MAE 69.259, RMSE 116.878, R² 0.490, and WAPE 54.239% on 379 observed jurisdiction-months.",
        "These values must be compared with the previous-month baseline rather than interpreted in isolation.",
    ])
    add_code(doc, "Code 6 — Compare the baseline and isolate failure cases", '''baseline_pred = test["lag1_covid"].to_numpy()
baseline_metrics = {
    "MAE": mean_absolute_error(actual, baseline_pred),
    "RMSE": mean_squared_error(actual, baseline_pred) ** 0.5,
    "R2": r2_score(actual, baseline_pred),
    "WAPE_pct": 100 * np.abs(actual - baseline_pred).sum() / actual.sum(),
}
errors = test[["state", "start_date", "covid_deaths"]].copy()
errors["prediction"] = pred
errors["absolute_error"] = np.abs(errors["covid_deaths"] - errors["prediction"])
print(baseline_metrics)
print(errors.nlargest(10, "absolute_error"))''', [
        "The baseline reaches MAE 48.892, RMSE 84.661, R² 0.732, and WAPE 38.289%, outperforming the model on every test metric.",
        "The gradient model's MAE is 41.66% worse than the baseline, so positive R² is not sufficient evidence for deployment.",
        "California January 2023 is the largest model miss: actual 1,327 versus predicted 2,550.8, an absolute error of 1,223.8 deaths.",
    ])

    section(doc, 6, "Results and Findings")
    add_table(doc, ["Model", "MAE", "RMSE", "R²", "WAPE"], [
        ["Gradient-boosted model", "69.259", "116.878", "0.490", "54.239%"],
        ["Previous-month baseline", "48.892", "84.661", "0.732", "38.289%"],
    ], "Table 2. Final January–August 2023 holdout performance; lower MAE/RMSE/WAPE and higher R² are better.")
    add_table(doc, ["Validation block", "Model MAE", "Baseline MAE", "Winner", "Interpretation"], [
        ["2021 H2", "359.550", "395.694", "Model", "Nonlinear lag structure helps during volatile waves."],
        ["2022 H1", "192.253", "386.325", "Model", "Model captures part of the winter-peak decline."],
        ["2022 H2", "118.633", "58.696", "Baseline", "Persistence dominates in a smoother low-burden regime."],
        ["2023 Jan–Aug test", "69.259", "48.892", "Baseline", "Late-period dynamics remain closer to persistence."],
    ], "Table 3. Expanding-window evidence shows that relative performance changes by pandemic regime.")
    para(doc, "MAE answers the operational question 'how many deaths off is the typical jurisdiction-month forecast?' RMSE is much larger than MAE for both methods, confirming a small number of large-state errors. R² remains positive for the gradient model, but the stronger baseline explains substantially more variance. WAPE shows the model's total absolute error equals 54.239% of total observed deaths in the test rows, which is too high for planning without wide uncertainty bands.")
    para(doc, "The failure is systematic rather than random noise. The model overpredicts several large jurisdictions during the 2023 decline, including California, Florida, and Texas. Tree boosting interpolates historical lag patterns but does not extrapolate a continuing structural decline well. The log target reduces peak dominance, yet the learned mapping still reflects the high-wave training regimes.")

    section(doc, 7, "Visualizations (described in text)")
    add_figure(doc, "w4_forecast_monthly.png", "Figure 1. Aggregate actual, gradient-boosted, and previous-month forecasts over the final eight-month test period.")
    para(doc, "Technical reading: lines sum predictions across available scored jurisdictions for each test month. Analytical reading: the boosted model remains too high through much of the decline, whereas the previous-month line follows persistence more closely. Aggregation can hide state errors, so this view is paired with point-level diagnostics.")
    add_figure(doc, "w4_actual_predicted.png", "Figure 2. Actual versus predicted jurisdiction-month counts on symmetric log-scaled axes; the dashed line is perfect prediction.")
    para(doc, "Technical reading: points above the diagonal are overpredictions and points below are underpredictions; the log-like scale shows both small and large states. Analytical reading: dispersion expands at high counts, confirming heteroscedasticity and explaining why RMSE is especially sensitive to large-state misses.")
    add_figure(doc, "w4_worst_errors.png", "Figure 3. Ten largest absolute errors in the January–August 2023 holdout.")
    para(doc, "Technical reading: bars rank state-months by absolute error in deaths. Analytical reading: California, Texas, Florida, and Arizona recur among the worst misses. A national-average score would understate the operational risk concentrated in populous jurisdictions.")

    section(doc, 8, "Challenges and How They Were Handled")
    bullets(doc, [
        "Temporal leakage was the largest modeling risk. I used shifted features, expanding validation, and a future holdout instead of random splitting.",
        "Suppressed counts are interval-censored. Midpoint 5 is used only for lag predictors; suppressed target rows are not evaluated as exact observations.",
        "Jurisdictions differ sharply in scale. The target uses log1p, numeric features are standardized, and evaluation returns to the original count scale.",
        "The final month is incomplete. September 2023 is excluded through the Week 1 partial-period flag.",
        "A positive model R² could look acceptable without context. A previous-month baseline and four metrics reveal that the model is inferior in 2023.",
        "Performance changes by pandemic regime. Multiple chronological folds expose the reversal between early 2022 and late 2022 rather than averaging it away.",
    ])

    section(doc, 9, "Discussion")
    para(doc, "The practical lesson is that model complexity must earn its place. The gradient model captures volatile transitions in two validation blocks, but late-2022 and 2023 behavior is closer to persistence. If deployed, its overprediction in large states could over-allocate scarce staff or supplies; underprediction during a new wave could be worse. A point forecast without prediction intervals is not enough for mortality planning.")
    para(doc, "The holdout also explains metric trade-offs. MAE is straightforward for routine planning, RMSE emphasizes rare large misses, R² supports comparison of explained variation, and WAPE is useful when volumes differ. None should be reported alone. MAPE was not used because small or suppressed targets make percentage errors unstable or undefined.")
    para(doc, "Limitations include a small time span, aggregate data, changing variants and reporting practices, missing predictors such as vaccination and hospital admissions, and a frozen source. HistGradientBoosting cannot extrapolate trends beyond combinations seen during training. The one-hot state effects are descriptive and may become stale if geography-specific dynamics shift.")
    para(doc, "The next improvement should retain the strong lag-1 forecast as an explicit skip path and train a model only to predict the correction. That reframes learning around month-to-month change, creates a safe zero-correction starting point, and is tested in Week 5 with a small regularized neural network.")

    section(doc, 10, "Conclusion")
    para(doc, "A leak-free gradient-boosted regression pipeline was trained on 1,690 jurisdiction-months and tested on 379 future rows. Although it achieved R² 0.490, its MAE of 69.259 and WAPE of 54.239% were worse than the previous-month baseline's MAE 48.892 and WAPE 38.289%. The model is therefore rejected for deployment in its current form.")
    para(doc, "Week 5 will use the same chronological split and features but redesign the target as a log-scale correction to the previous month. The goal is not to add depth for its own sake; it is to preserve baseline behavior unless the network finds evidence for a better adjustment.")
    section(doc, 11, "References"); references(doc)
    return save_doc(doc, "Week_4_Supervised_Learning_Model.docx")


def build_week5():
    d = S["deep_learning"]
    doc = Document(); configure_document(doc, 5)
    add_title_page(doc, "Deep Learning Application in Data Science", 5, "A regularized residual neural network for monthly mortality forecasting")
    add_report_map(doc)
    section(doc, 2, "Introduction")
    para(doc, "This report redesigns the Week 4 forecast as a residual neural-network problem. Instead of asking a network to learn the full count from scratch, it predicts a log-scale correction to the previous month's count. A zero correction exactly reproduces the strong persistence baseline, while positive and negative corrections represent expected growth and decline.")
    para(doc, "The CDC dataset remains appropriate because its 52 jurisdiction time series provide nonlinear lag, seasonal, and geographic relationships. It is also deliberately challenging for deep learning: only 2,069 feature-ready rows exist, so architecture size, regularization, and chronological validation matter more than adding layers.")
    callout(doc, "Main outcome.", "The residual network improves the January–August 2023 test MAE from 48.892 to 34.253 deaths and raises R² from 0.732 to 0.867. The improvement comes from architecture aligned with the baseline, not from a large network.", fill="E7F4EC")
    section(doc, 3, "Dataset Overview")
    para(doc, "Input preparation is carried forward from Week 4: All Sexes, All Ages, monthly records for 52 non-overlapping jurisdictions; September 2023 is excluded; three COVID lags, a three-month average, lagged all-cause and pneumonia counts, seasonal sine/cosine, time index, and one-hot state identity produce 61 input features.")
    add_table(doc, ["Partition", "Rows", "Calendar period", "Role"], [
        ["Fit", "1,381", "Apr 2020–Jun 2022", "Gradient updates."],
        ["Validation", "309", "Jul–Dec 2022", "Early stopping and learning-rate decisions."],
        ["Test", "379", "Jan–Aug 2023", "One-time future evaluation."],
    ], "Table 1. Chronological deep-learning partitions.")
    para(doc, "The supervised target is log1p(current COVID deaths) minus log1p(previous-month COVID deaths). This residual is centered closer to zero and converts multiplicative changes into additive corrections. Evaluation is performed after adding the predicted residual to the log baseline and transforming back to death counts.")

    section(doc, 4, "Methodology", "The method defines the forecasting task, justifies the network architecture and hyperparameters, trains with explicit overfitting controls, evaluates four metrics, and diagnoses where the network still fails.")
    subheading(doc, "Step 1 — Preserve the Week 4 leak-free features")
    para(doc, "No current-month mortality measure is an input. Numeric preprocessing is fit on the 2020–June 2022 fit set only, and state categories are one-hot encoded with unknown-category handling. The validation and test arrays are transformed with the fitted training parameters.")
    subheading(doc, "Step 2 — Use a residual target and skip-path logic")
    para(doc, "The previous-month count is a strong forecast in smooth periods. The output layer is initialized to zero, so the untrained model begins at that baseline after the skip addition. This is safer than a random full-count prediction and focuses capacity on changes the baseline misses.")
    subheading(doc, "Step 3 — Choose a compact architecture")
    para(doc, "The network has Dense(32) → Dropout(20%) → Dense(16) → Dropout(15%) → Dense(1), with ReLU hidden activations and a linear correction output. The taper from 32 to 16 compresses interactions among 61 inputs without creating a large parameter count; the complete model has only 2,529 trainable parameters.")
    subheading(doc, "Why ReLU, linear output, and Huber loss")
    para(doc, "ReLU keeps positive hidden activations non-saturating and is a better fit than sigmoid here because the hidden representation is not a probability and standardized inputs can span both signs. Tanh could represent signed activations but saturates at large magnitude. A linear final unit is required because the correction can be positive or negative. Huber loss with delta 0.3 behaves quadratically near small residual errors but linearly for large wave shocks, reducing the influence of extreme transitions compared with mean squared error.")
    subheading(doc, "Step 4 — Address overfitting directly")
    para(doc, "L2 weight penalty 0.001 discourages large coefficients, dropout randomly removes 20% then 15% of hidden activations during training, and early stopping restores the lowest validation-loss weights after 20 non-improving epochs. Adam uses learning rate 0.001, batch size 64, and ReduceLROnPlateau. The regularized network's best epoch is 4 and training stops at 24.")
    subheading(doc, "Step 5 — Create an unregularized reference")
    para(doc, "I trained the same dense widths for 160 epochs without dropout or L2. Its best validation loss occurs at epoch 1; by the final epoch training loss is 0.01497 while validation loss is 0.49378, a gap of 0.47881. This is direct evidence that the dataset is too small for unconstrained fitting.")
    subheading(doc, "Step 6 — Evaluate against the same baseline")
    para(doc, "The final 2023 test uses MAE, RMSE, R², and WAPE on original death counts. The comparison baseline is exactly the previous-month forecast embedded in the residual architecture. Worst state-month errors are reported so an improved average cannot hide specific failures.")

    section(doc, 5, "Python Code Sections")
    add_code(doc, "Code 1 — Imports, seeds, and chronological arrays", '''import os
import random

import numpy as np
import pandas as pd
import tensorflow as tf
from sklearn.compose import ColumnTransformer
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.preprocessing import OneHotEncoder, StandardScaler

SEED = 42
os.environ["TF_DETERMINISTIC_OPS"] = "1"
random.seed(SEED)
np.random.seed(SEED)
tf.keras.utils.set_random_seed(SEED)

# model_data and the real Week 4 feature names are created exactly as in Week 4.
fit_data = model_data[model_data["start_date"] < "2022-07-01"]
val_data = model_data[model_data["start_date"].between("2022-07-01", "2022-12-01")]
test_data = model_data[model_data["start_date"].between("2023-01-01", "2023-08-01")]''', [
        "Deterministic seeds make weight initialization, dropout masks, and reported results repeatable within the same software/hardware setup.",
        "The split is chronological rather than random: 1,381 fit rows, 309 validation rows, and 379 future test rows.",
        "The test period is not used for architecture, stopping, or learning-rate decisions.",
    ])
    add_code(doc, "Code 2 — Fit preprocessing only on the fit period", '''num_features = [
    "lag1_covid", "lag2_covid", "lag3_covid", "rolling3_covid",
    "lag1_total", "lag1_pneumonia", "month_sin", "month_cos", "time_index"
]
all_features = num_features + ["state"]
preprocess = ColumnTransformer([
    ("numeric", StandardScaler(), num_features),
    ("state", OneHotEncoder(handle_unknown="ignore", sparse_output=False), ["state"]),
])

X_fit = preprocess.fit_transform(fit_data[all_features]).astype("float32")
X_val = preprocess.transform(val_data[all_features]).astype("float32")
X_test = preprocess.transform(test_data[all_features]).astype("float32")
print(X_fit.shape, X_val.shape, X_test.shape)''', [
        "Fitting StandardScaler and the one-hot vocabulary only on the fit period prevents validation/test distribution information from entering preprocessing.",
        "Each transformed row has 61 inputs: nine numeric features and 52 jurisdiction indicators.",
        "Float32 reduces memory and matches TensorFlow's default numerical type without losing meaningful precision for these features.",
    ])
    add_code(doc, "Code 3 — Define the log-residual target", '''y_fit = (
    np.log1p(fit_data["covid_deaths"].to_numpy())
    - np.log1p(fit_data["lag1_covid"].to_numpy())
).astype("float32")
y_val = (
    np.log1p(val_data["covid_deaths"].to_numpy())
    - np.log1p(val_data["lag1_covid"].to_numpy())
).astype("float32")
y_test = test_data["covid_deaths"].to_numpy(dtype="float32")''', [
        "A residual of zero means no change from the previous month; negative and positive values represent proportional decline and growth.",
        "log1p supports zero counts and reduces the difference between small and large jurisdictions.",
        "The final y_test remains on the original count scale because all reported metrics must have operational meaning.",
    ])
    add_code(doc, "Code 4 — Build the compact regularized residual network", '''regularizer = tf.keras.regularizers.l2(1e-3)
model = tf.keras.Sequential([
    tf.keras.layers.Input(shape=(61,)),
    tf.keras.layers.Dense(32, activation="relu", kernel_regularizer=regularizer),
    tf.keras.layers.Dropout(0.20),
    tf.keras.layers.Dense(16, activation="relu", kernel_regularizer=regularizer),
    tf.keras.layers.Dropout(0.15),
    tf.keras.layers.Dense(
        1, activation="linear", kernel_initializer="zeros", bias_initializer="zeros"
    ),
])
model.compile(
    optimizer=tf.keras.optimizers.Adam(learning_rate=1e-3),
    loss=tf.keras.losses.Huber(delta=0.3),
    metrics=[tf.keras.metrics.MeanAbsoluteError(name="mae_log_residual")],
)
print(model.count_params())''', [
        "The model contains 2,529 parameters, intentionally small for 1,381 fit rows.",
        "ReLU captures nonlinear lag interactions without the saturation of sigmoid; the linear output permits both upward and downward corrections.",
        "Zero output initialization makes the first forecast equal to the previous-month baseline before any correction is learned.",
    ])
    add_code(doc, "Code 5 — Train with early stopping and learning-rate reduction", '''callbacks = [
    tf.keras.callbacks.EarlyStopping(
        monitor="val_loss", patience=20, min_delta=1e-5,
        restore_best_weights=True,
    ),
    tf.keras.callbacks.ReduceLROnPlateau(
        monitor="val_loss", patience=10, factor=0.5, min_lr=1e-5,
    ),
]
history = model.fit(
    X_fit, y_fit,
    validation_data=(X_val, y_val),
    epochs=250, batch_size=64, shuffle=True, verbose=0,
    callbacks=callbacks,
)
best_epoch = int(np.argmin(history.history["val_loss"]) + 1)
print(best_epoch, len(history.history["loss"]))''', [
        "Validation loss is lowest at epoch 4, and training stops at epoch 24 after patience is exhausted.",
        "restore_best_weights ensures test predictions use epoch 4 rather than the overfit final epoch.",
        "ReduceLROnPlateau allows smaller adjustments if validation improvement stalls before early stopping.",
    ])
    add_code(doc, "Code 6 — Reconstruct count forecasts and compute four metrics", '''log_correction = model.predict(X_test, verbose=0).ravel()
baseline = test_data["lag1_covid"].to_numpy()
pred = np.clip(
    np.expm1(np.log1p(baseline) + log_correction), 0, None
)

metrics = {
    "MAE": mean_absolute_error(y_test, pred),
    "RMSE": mean_squared_error(y_test, pred) ** 0.5,
    "R2": r2_score(y_test, pred),
    "WAPE_pct": 100 * np.abs(y_test - pred).sum() / y_test.sum(),
}
print(metrics)''', [
        "The predicted correction is added to the log previous-month count, implementing the explicit skip path.",
        "On the future test, the network reaches MAE 34.253, RMSE 59.621, R² 0.867, and WAPE 26.825%.",
        "All metrics are calculated after inverse transformation, so errors are reported in deaths rather than abstract normalized units.",
    ])
    add_code(doc, "Code 7 — Train an unregularized reference for evidence of overfitting", '''unregularized = tf.keras.Sequential([
    tf.keras.layers.Input(shape=(61,)),
    tf.keras.layers.Dense(32, activation="relu"),
    tf.keras.layers.Dense(16, activation="relu"),
    tf.keras.layers.Dense(1, kernel_initializer="zeros", bias_initializer="zeros"),
])
unregularized.compile(
    optimizer=tf.keras.optimizers.Adam(1e-3),
    loss=tf.keras.losses.Huber(delta=0.3),
)
ref_history = unregularized.fit(
    X_fit, y_fit, validation_data=(X_val, y_val),
    epochs=160, batch_size=64, verbose=0,
)
print(ref_history.history["loss"][-1], ref_history.history["val_loss"][-1])''', [
        "The reference removes dropout, L2, and early stopping while retaining the same dense widths and residual target.",
        "Its final training loss falls to 0.01497, but validation loss rises to 0.49378; lower training loss is not better generalization.",
        "The 0.47881 final gap and epoch-1 validation optimum provide direct evidence for keeping regularization and early stopping.",
    ])

    section(doc, 6, "Results and Findings")
    add_table(doc, ["Method", "MAE", "RMSE", "R²", "WAPE"], [
        ["Residual neural network", "34.253", "59.621", "0.867", "26.825%"],
        ["Previous-month baseline", "48.892", "84.661", "0.732", "38.289%"],
        ["Week 4 boosted model", "69.259", "116.878", "0.490", "54.239%"],
    ], "Table 2. Same January–August 2023 holdout; lower error and higher R² are better.")
    para(doc, "The residual network reduces MAE by 29.94% and RMSE by 29.58% relative to the previous-month baseline. WAPE falls by 11.464 percentage points, and R² increases by 0.135. It also reverses the Week 4 model failure because its architecture cannot forget the baseline unless the learned correction moves it.")
    add_table(doc, ["Largest remaining error", "Actual", "Prediction", "Absolute error"], [
        ["Texas, Feb 2023", "632", "965.2", "333.2"],
        ["Florida, Jan 2023", "991", "685.2", "305.8"],
        ["North Carolina, Jan 2023", "675", "372.7", "302.3"],
        ["Texas, Jan 2023", "1,109", "835.6", "273.4"],
        ["California, Feb 2023", "754", "1,026.5", "272.5"],
    ], "Table 3. The model improves average performance but still misses rapid transitions in large jurisdictions.")
    para(doc, "The error list is balanced between over- and underprediction around the January transition. This matters operationally: a smaller overall MAE does not remove surge risk. The model needs intervals, rolling recalibration, and external leading indicators before use for decisions.")

    section(doc, 7, "Visualizations (described in text)")
    add_figure(doc, "w5_architecture.png", "Figure 1. Compact residual network. The learned log correction is added to log1p(previous-month deaths) before inverse transformation.")
    para(doc, "Technical reading: 61 inputs flow through 32 and 16 ReLU units with dropout, then one linear correction. Analytical reading: the architecture allocates most responsibility to the proven previous-month estimate and asks the network only to model departures. This design is better matched to a small, persistent time series than a large unconstrained multilayer network.")
    add_figure(doc, "w5_training_curves.png", "Figure 2. Regularized network training and validation Huber loss; best validation epoch is 4 and restored after stopping at epoch 24.")
    para(doc, "Technical reading: separate curves show optimization fit and future-period validation loss by epoch. Analytical reading: validation improvement occurs early; continued fitting mainly reduces training error. Early stopping prevents the later gap from being converted into test predictions.")
    add_figure(doc, "w5_overfitting_reference.png", "Figure 3. Unregularized reference training and validation loss over 160 epochs.")
    para(doc, "Technical reading: training loss continues downward while validation loss is lowest at epoch 1 and ends at 0.49378. Analytical reading: the network memorizes historical jurisdiction-wave combinations that do not transfer to late 2022. The divergence is evidence, not a generic warning, for regularization.")
    add_figure(doc, "w5_forecast_monthly.png", "Figure 4. Aggregate neural, baseline, and actual test-period counts across scored jurisdictions.")
    para(doc, "Technical reading: monthly lines aggregate the same 379 observed state-month rows used for metrics. Analytical reading: the residual correction tracks the broad 2023 decline more closely than the lag baseline. The remaining January–February gaps align with the largest state-level errors.")

    section(doc, 8, "Challenges and How They Were Handled")
    bullets(doc, [
        "The dataset is small for deep learning. I used two hidden layers and 2,529 parameters instead of a wide/deep architecture.",
        "A full-count network initially tended to underfit or overfit changing regimes. The target was redesigned as a log residual around the strong lag-1 baseline.",
        "Overfitting appeared quickly. Dropout, L2, chronological validation, early stopping, and a reference curve make the response measurable.",
        "Counts are extremely skewed across states and waves. log1p residuals, Huber loss, and original-scale evaluation reduce domination without hiding errors.",
        "CPU resources limit exhaustive tuning. A compact deterministic run and narrow, justified hyperparameter set are more reproducible than an unreported search.",
        "The model still makes large transition errors. Worst cases are listed, and no claim of deployment readiness is made without uncertainty intervals and live data.",
    ])

    section(doc, 9, "Discussion")
    para(doc, "The architecture choice matters more than the label 'deep learning.' A generic network trained on full log counts performed poorly because the small sample spans major nonstationary regimes. The residual design begins from a useful forecast and learns modest changes. Zero initialization of the correction output is a practical safeguard: if training learns nothing, the model behaves like the baseline instead of emitting arbitrary counts.")
    para(doc, "ReLU is suitable in the hidden layers because the inputs are standardized and the objective needs sparse nonlinear interactions, not bounded probabilities. A sigmoid output would be wrong because corrections must be negative, while a linear output is appropriate. Huber loss protects training from wave shocks, but RMSE in evaluation remains important because planners still care about large misses.")
    para(doc, "Limitations remain: aggregate provisional data, no population or vaccination variables, only 44 complete months, interval-censored small counts, and a static source ending in 2023. TensorFlow randomness is controlled as far as practical, but exact floating-point results may vary across hardware. The model has no calibrated prediction intervals and may fail under a new variant or coding change.")
    para(doc, "If overfitting controls were skipped, the unregularized final validation loss shows what would happen: training performance would look excellent while future performance deteriorated. If the baseline were omitted, the network's positive R² could be celebrated without knowing whether it adds operational value. Both checks are essential.")

    section(doc, 10, "Conclusion")
    para(doc, "A 2,529-parameter TensorFlow residual network improved the future holdout to MAE 34.253, RMSE 59.621, R² 0.867, and WAPE 26.825%, outperforming the previous-month baseline and the Week 4 boosted model. Training curves show the best validation result at epoch 4 and clear unregularized overfitting by epoch 160.")
    para(doc, "Week 6 will integrate the cleaning, EDA, clustering, boosted-model failure, and residual-network improvement into one pipeline. The recommended forecasting system will keep the lag baseline as a fallback, monitor rolling WAPE, and update from a current CDC source rather than the frozen 2023 extract.")
    section(doc, 11, "References"); references(doc, include_tf=True)
    return save_doc(doc, "Week_5_Deep_Learning_Application.docx")


def build_week6():
    e, c, m, d = S["eda"], S["clustering"], S["supervised"], S["deep_learning"]
    doc = Document(); configure_document(doc, 6)
    add_title_page(doc, "Integrative Capstone Project and Evaluation", 6, "From censored CDC mortality data to profiles and one-month forecasts")
    add_report_map(doc)
    section(doc, 2, "Introduction")
    para(doc, "This capstone combines the six-week workflow into one reproducible public-health data science pipeline. The project acquires and audits the CDC Provisional COVID-19 Deaths by Sex and Age extract, performs grain-safe exploratory analysis, clusters jurisdictions by mortality profile, evaluates a conventional supervised model honestly, and builds a compact residual neural network that improves a strong operational baseline.")
    para(doc, "The project opportunity is twofold: describe how burden differs by time, age, sex, respiratory overlap, and jurisdiction; then test whether recent history can support one-month-ahead jurisdiction forecasts. The dataset was selected because its privacy suppression, overlapping aggregates, right-skewed counts, and nonstationary waves require connected decisions across the full pipeline rather than isolated code demonstrations.")
    callout(doc, "Capstone result.", "Cleaning preserved all 137,700 records without zero-filling suppressed cells; EDA identified a January 2021 peak and 75.76% age-65+ share; four modest clusters described jurisdiction profiles; the Week 4 booster failed against persistence; a baseline-anchored residual network improved test MAE by 29.94%.")
    section(doc, 3, "Dataset Overview")
    para(doc, "The CDC/NCHS CSV contains 137,700 rows, 16 raw columns, 54 geography labels, three sex categories, 17 age groups, and monthly/yearly/cumulative grains. It covers 1 January 2020 through 23 September 2023 and was published as of 27 September 2023. The publisher states that the table is no longer updated; the capstone therefore demonstrates method on a historical extract.")
    add_table(doc, ["Quality characteristic", "Observed evidence", "Pipeline response"], [
        ["Privacy suppression", "39,430 COVID; 44,864 pneumonia blanks; all on footnoted rows", "Retain NaN; add measure flags; midpoint only for lag predictors."],
        ["Mixed grains", "123,930 monthly; 11,016 yearly; 2,754 total", "Require explicit Group filter."],
        ["Overlapping age schemes", "17 groups include incompatible bands", "Use 11 mutually exclusive bands for sums."],
        ["Partial endpoint", "5,508 Sep/2023 rows", "Flag; exclude from complete-period models."],
        ["Right-skewed counts", "CA, FL, TX exceed IQR cutoff", "Keep valid records; use shares/log residuals."],
        ["Geographic overlap", "US total and NYC nested in NY", "Exclude both from 52-unit modeling matrix."],
    ], "Table 1. Source characteristics and the connected response across project phases.")
    para(doc, "The clean master table has 25 columns after type conversion, six suppression flags, a row suppression flag, a partial-period marker, and a safe COVID share. It retains all source rows so later analytical slices can be rebuilt without returning to a modified raw file.")

    section(doc, 4, "Methodology", "This section condenses every phase from Weeks 1–5 in the order of a production data science pipeline, while retaining the rationale behind each decision.")
    subheading(doc, "Phase 1 — Acquisition and reproducibility")
    para(doc, "The attached file was matched to CDC Socrata identifier 9bhg-hcku by title, schema, publisher, and update date. The raw input is treated as immutable. Paths, exact source URL, software libraries, deterministic seeds, and fixed split dates are documented so the analysis can be rerun.")
    subheading(doc, "Phase 2 — Cleaning and preprocessing")
    para(doc, "All fields were first loaded as text. Dates use explicit MM/DD/YYYY parsing; commas are removed from six documented count fields; Year/Month use nullable integers. Suppressed 1–9 cells stay missing and receive measure-level flags. Domain assertions check duplicates, dates, nonnegativity, totals, and overlap logic. September 2023 and 2023 yearly records are marked partial.")
    subheading(doc, "Phase 3 — Exploratory analysis")
    para(doc, "National monthly trends use By Month/United States/All Sexes/All Ages. Age and sex comparisons use the 11-band exclusive scheme. Jurisdiction comparison uses 52 non-overlapping units and COVID-19 as a percentage of all deaths. Spearman correlation summarizes aligned complete-month mortality series without assuming normality.")
    subheading(doc, "Phase 4 — Unsupervised profiles")
    para(doc, "Seven percentages—COVID, pneumonia, overlap, influenza, age 65+, male share, and 2021 timing—are standardized. K-Means candidates k=2–8 are compared with inertia and silhouette; k=4 has the highest silhouette at 0.222. PCA and Ward linkage visualize the groups without replacing the seven-dimensional fit.")
    subheading(doc, "Phase 5 — Conventional supervised benchmark")
    para(doc, "The one-month regression uses three COVID lags, rolling mean, lagged total/pneumonia, seasonal encoding, time trend, and one-hot state. HistGradientBoosting is trained on log counts with chronological validation. The previous-month forecast is the baseline. On 2023, the booster is rejected because all four metrics are worse than persistence.")
    subheading(doc, "Phase 6 — Deep residual model")
    para(doc, "A compact 32→16 ReLU network predicts the log correction to the previous month. Zero output initialization preserves the baseline at start; dropout, L2, Huber loss, and early stopping address the small sample. The fit/validation/test windows are chronological, and an unregularized reference provides direct overfitting evidence.")
    subheading(doc, "Phase 7 — Evaluation and decision rules")
    para(doc, "Model selection uses future-block MAE, RMSE, R², WAPE, baseline comparison, and worst-case inspection. Descriptive results are checked against exact tables. No model is recommended because it is complex or has positive R²; it must improve the baseline and remain stable under rolling monitoring.")

    section(doc, 5, "Python Code Sections")
    add_code(doc, "Code 1 — Acquire, parse, and preserve suppression", '''from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from sklearn.cluster import KMeans
from sklearn.compose import ColumnTransformer
from sklearn.decomposition import PCA
from sklearn.ensemble import HistGradientBoostingRegressor
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score, silhouette_score
from sklearn.preprocessing import OneHotEncoder, StandardScaler
import tensorflow as tf

raw = pd.read_csv("Provisional COVID-19 Deaths by Sex and Age.csv",
                  dtype=str, low_memory=False)
df = raw.rename(columns={
    "Data As Of": "data_as_of", "Start Date": "start_date", "End Date": "end_date",
    "Group": "group", "Year": "year", "Month": "month", "State": "state",
    "Sex": "sex", "Age Group": "age_group", "COVID-19 Deaths": "covid_deaths",
    "Total Deaths": "total_deaths", "Pneumonia Deaths": "pneumonia_deaths",
    "Pneumonia and COVID-19 Deaths": "pneumonia_covid_deaths",
    "Influenza Deaths": "influenza_deaths",
    "Pneumonia, Influenza, or COVID-19 Deaths": "pic_deaths", "Footnote": "footnote",
})''', [
        "The capstone begins from the actual attached filename and maps every source field used later to a consistent name.",
        "Loading as text prevents chunk-based mixed typing before count and date rules are applied.",
        "The raw shape is 137,700 × 16, providing a fixed acquisition checkpoint.",
    ])
    add_code(doc, "Code 2 — Enforce quality rules and partial-period metadata", '''count_cols = ["covid_deaths", "total_deaths", "pneumonia_deaths",
              "pneumonia_covid_deaths", "influenza_deaths", "pic_deaths"]
for col in ["data_as_of", "start_date", "end_date"]:
    df[col] = pd.to_datetime(df[col], format="%m/%d/%Y", errors="raise")
for col in ["year", "month"]:
    df[col] = pd.to_numeric(df[col], errors="coerce").astype("Int64")
for col in count_cols:
    df[col] = pd.to_numeric(df[col].str.replace(",", regex=False), errors="coerce")
    df[f"{col}_suppressed"] = df[col].isna() & df["footnote"].notna()

month_end = df["start_date"] + pd.offsets.MonthEnd(0)
year_end = pd.to_datetime(df["year"].astype("string") + "-12-31",
                          format="%Y-%m-%d", errors="coerce")
df["is_partial_period"] = (
    (df["group"].eq("By Month") & (df["end_date"] < month_end)) |
    (df["group"].eq("By Year") & (df["end_date"] < year_end))
)
assert not df.duplicated(["group", "year", "month", "state", "sex", "age_group"]).any()
assert not (df["covid_deaths"] > df["total_deaths"]).any()''', [
        "Suppression remains explicit rather than being changed to zero, preserving the CDC's confidentiality interval.",
        "The partial-period rule identifies 5,508 rows and is reused by EDA and modeling.",
        "Executable grain and domain assertions protect every downstream result from silent structural errors.",
    ])
    add_code(doc, "Code 3 — Produce the national EDA slices", '''exclusive_age = [
    "Under 1 year", "1-4 years", "5-14 years", "15-24 years", "25-34 years",
    "35-44 years", "45-54 years", "55-64 years", "65-74 years",
    "75-84 years", "85 years and over",
]
monthly_us = df.query(
    "group == 'By Month' and state == 'United States' "
    "and sex == 'All Sexes' and age_group == 'All Ages'"
)
complete_us = monthly_us.loc[~monthly_us["is_partial_period"]]
age_us = df.query(
    "group == 'By Total' and state == 'United States' and sex == 'All Sexes'"
).loc[lambda x: x["age_group"].isin(exclusive_age)]

peak = complete_us.loc[complete_us["covid_deaths"].idxmax()]
age65 = age_us.loc[age_us["age_group"].isin(
    ["65-74 years", "75-84 years", "85 years and over"]
), "covid_deaths"].sum()
print(peak["start_date"], peak["covid_deaths"])
print(100 * age65 / age_us["covid_deaths"].sum())''', [
        "The filters lock one grain and avoid adding aggregate rows to components.",
        "The national peak is January 2021 with 105,565 deaths; ages 65+ account for 75.76% of cumulative deaths.",
        "These exact checkpoints anchor the visualization narrative and later recommendations.",
    ])
    add_code(doc, "Code 4 — Select and fit the four-profile clustering", '''# features is the 52 × 7 percentage matrix documented in Week 3.
X_cluster = StandardScaler().fit_transform(features)
selection = []
for k in range(2, 9):
    candidate = KMeans(n_clusters=k, n_init=50, random_state=42)
    labels = candidate.fit_predict(X_cluster)
    selection.append((k, candidate.inertia_, silhouette_score(X_cluster, labels)))

cluster_model = KMeans(n_clusters=4, n_init=100, random_state=42)
cluster_labels = cluster_model.fit_predict(X_cluster)
pca_points = PCA(n_components=2).fit_transform(X_cluster)
print(selection)''', [
        "The feature matrix uses shares rather than raw counts, so distance represents mortality profile instead of population size.",
        "Four clusters are selected because k=4 has the highest tested silhouette (0.222), while the close scores are retained as a limitation.",
        "PCA is a display tool; assignments use all seven standardized features.",
    ])
    add_code(doc, "Code 5 — Build leak-free one-month forecasting features", '''panel = df.query(
    "group == 'By Month' and sex == 'All Sexes' and age_group == 'All Ages'"
).loc[lambda x: ~x["state"].isin(["United States", "New York City"])
 & ~x["is_partial_period"]].sort_values(["state", "start_date"])

for col in ["covid_deaths", "total_deaths", "pneumonia_deaths"]:
    panel[f"{col}_for_lag"] = panel[col].fillna(5.0)
g = panel.groupby("state")
for lag in [1, 2, 3]:
    panel[f"lag{lag}_covid"] = g["covid_deaths_for_lag"].shift(lag)
panel["rolling3_covid"] = g["covid_deaths_for_lag"].transform(
    lambda s: s.shift(1).rolling(3).mean()
)
panel["lag1_total"] = g["total_deaths_for_lag"].shift(1)
panel["lag1_pneumonia"] = g["pneumonia_deaths_for_lag"].shift(1)''', [
        "All count predictors are shifted within state, ensuring the current target is unavailable at forecast time.",
        "A midpoint is used only for censored lag inputs; suppressed target rows are excluded from metric calculations.",
        "The resulting 2,069-row panel supports identical chronological comparisons across conventional and neural models.",
    ])
    add_code(doc, "Code 6 — Define the residual neural correction", '''regularizer = tf.keras.regularizers.l2(1e-3)
network = tf.keras.Sequential([
    tf.keras.layers.Input(shape=(61,)),
    tf.keras.layers.Dense(32, activation="relu", kernel_regularizer=regularizer),
    tf.keras.layers.Dropout(0.20),
    tf.keras.layers.Dense(16, activation="relu", kernel_regularizer=regularizer),
    tf.keras.layers.Dropout(0.15),
    tf.keras.layers.Dense(1, activation="linear",
                          kernel_initializer="zeros", bias_initializer="zeros"),
])
network.compile(optimizer=tf.keras.optimizers.Adam(1e-3),
                loss=tf.keras.losses.Huber(delta=0.3))
# y = log1p(current deaths) - log1p(previous-month deaths)
# final forecast = expm1(log1p(previous month) + predicted correction)''', [
        "The 2,529-parameter network is deliberately compact for 1,381 fit rows.",
        "The zero-initialized linear output starts from the lag-1 forecast and permits signed corrections.",
        "Dropout, L2, Huber loss, and early stopping respond to the overfitting seen in the unregularized reference.",
    ])
    add_code(doc, "Code 7 — Apply a baseline-first model acceptance rule", '''def score(actual, prediction):
    return {
        "MAE": mean_absolute_error(actual, prediction),
        "RMSE": mean_squared_error(actual, prediction) ** 0.5,
        "R2": r2_score(actual, prediction),
        "WAPE_pct": 100 * np.abs(actual - prediction).sum() / actual.sum(),
    }

baseline_score = score(y_test, lag1_prediction)
boosted_score = score(y_test, boosted_prediction)
neural_score = score(y_test, neural_prediction)
accepted = (
    neural_score["MAE"] < baseline_score["MAE"]
    and neural_score["RMSE"] < baseline_score["RMSE"]
    and neural_score["WAPE_pct"] < baseline_score["WAPE_pct"]
)
print(baseline_score, boosted_score, neural_score, accepted)''', [
        "The same observed 2023 target rows are scored for every method, making the comparison fair.",
        "The booster fails the acceptance rule; the residual network passes with MAE 34.253, RMSE 59.621, R² 0.867, and WAPE 26.825%.",
        "Acceptance is conditional on this historical holdout; live deployment still requires intervals, drift monitoring, and current data.",
    ])

    section(doc, 6, "Results and Findings")
    add_table(doc, ["Pipeline phase", "Result", "Meaning"], [
        ["Cleaning", "137,700 rows retained; 0 key duplicates; 5,508 partial rows flagged", "Representation improved without deleting valid events."],
        ["Suppression", "COVID blanks 28.63%; pneumonia blanks 32.58%", "Censoring is material and non-random."],
        ["National EDA", "1,146,774 COVID deaths; 9.32% of 12,303,399 total", "Historical cumulative burden in source period."],
        ["Peak month", "Jan 2021: 105,565; 28.25% of all deaths", "Wave structure dominates averages."],
        ["Age", "65+: 868,831; 75.76%", "Burden is concentrated among older adults."],
        ["Sex", "Male/female 1.218; max ratio 1.75 at 45–54", "Midlife proportional gap is visible beyond oldest-age counts."],
        ["Clustering", "k=4; silhouette 0.222; sizes 8/13/23/8", "Useful but soft profile segmentation."],
        ["Boosted test", "MAE 69.259; R² 0.490; WAPE 54.239%", "Worse than persistence; rejected."],
        ["Residual neural test", "MAE 34.253; R² 0.867; WAPE 26.825%", "29.94% MAE improvement over persistence."],
    ], "Table 2. Integrated capstone outcomes.")
    para(doc, "The strongest project finding is methodological: preserving the baseline within the architecture mattered more than increasing model complexity. The Week 4 tree model had a respectable positive R² but failed every comparison against persistence. The residual network improved because it learned a bounded correction around a credible starting point and stopped after early validation gains.")
    para(doc, "Descriptively, mortality burden is concentrated by age and wave but varies meaningfully in composition. Four clusters capture older-age concentration, male skew, pneumonia overlap, and year timing. Their modest separation means they are best used to choose peer comparisons and hypotheses, not to declare one group successful or unsuccessful.")

    section(doc, 7, "Visualizations (described in text)")
    add_figure(doc, "w2_national_monthly.png", "Figure 1. Complete-period national monthly COVID-19 trend with partial September 2023 shaded.")
    para(doc, "Technical reading: the x-axis is month and y-axis is COVID-19 deaths, with January 2021 annotated. Analytical reading: repeated waves and a large 2021 peak make random validation and simple averages unsafe. The partial-period shading shows how data-quality metadata changes interpretation.")
    add_figure(doc, "w2_age_profile.png", "Figure 2. Cumulative count and within-age death-share across mutually exclusive age groups.")
    para(doc, "Technical reading: bars and a secondary-axis line represent count and percentage. Analytical reading: the oldest group has the largest count, while 65–74 has the highest COVID share of all deaths. Policy interpretation changes depending on which measure is used.")
    add_figure(doc, "w3_cluster_heatmap.png", "Figure 3. Standardized four-cluster centroid profiles across seven engineered percentages.")
    para(doc, "Technical reading: values are standard deviations from the 52-jurisdiction mean. Analytical reading: Cluster 1's older-age concentration and Cluster 3's pneumonia overlap/younger male profile are multivariate differences that a raw-count ranking would miss.")
    add_figure(doc, "w4_forecast_monthly.png", "Figure 4. Week 4 boosted model versus persistence on the future holdout.")
    para(doc, "Technical reading: aggregated forecast lines use the same observed state-months. Analytical reading: the booster stays too high through decline, explaining why it loses despite positive R². This is the rejected benchmark that motivates residual design.")
    add_figure(doc, "w5_training_curves.png", "Figure 5. Residual neural-network learning curves with the best validation epoch marked.")
    para(doc, "Technical reading: Huber loss is plotted for fit and validation sets by epoch. Analytical reading: useful learning happens immediately and additional fitting risks memorization. Restoring epoch 4 is justified by observed validation behavior.")
    add_figure(doc, "w5_forecast_monthly.png", "Figure 6. Residual neural forecast, persistence baseline, and actual 2023 holdout totals.")
    para(doc, "Technical reading: all lines aggregate identical scored rows. Analytical reading: learned corrections improve the broad decline while large January–February transition errors remain. The model should be a monitored challenger, not an unqualified replacement.")

    section(doc, 8, "Challenges and How They Were Handled")
    add_table(doc, ["Challenge", "Resolution", "Evidence"], [
        ["Censored small counts", "Preserve NaN and flags; midpoint only for lag inputs", "All missing count cells align with privacy-footnoted rows."],
        ["Double-counting risk", "Explicit Group/Sex/Age filters and exclusive age bands", "Exclusive bands reconcile to All Ages."],
        ["Valid extreme counts", "Retain CA/FL/TX; use shares/log residuals", "Zero domain violations despite IQR flags."],
        ["Weak cluster boundaries", "Report silhouette and soft interpretation", "Best 0.222; k=7 close at 0.220."],
        ["Temporal regime change", "Expanding windows and future holdout", "Boosted model reverses from early-fold wins to late losses."],
        ["Neural overfitting", "Compact model, dropout, L2, Huber, early stopping", "Unregularized validation loss 0.49378 vs train 0.01497."],
        ["Frozen data source", "Limit claims to historical period; recommend CDC WONDER/current feeds", "Publisher stopped updates after 27 Sep 2023."],
    ], "Table 3. Project challenges, actions, and observed evidence.")

    section(doc, 9, "Discussion")
    para(doc, "The pipeline demonstrates that cleaning decisions determine modeling validity. Suppression cannot be treated as zero, overlapping categories cannot be freely summed, and partial periods cannot be model targets. Each of those mistakes would improve apparent completeness while degrading truth. Retaining flags allows the same clean master table to serve descriptive and predictive tasks with different, explicit assumptions.")
    para(doc, "The models also show why evaluation must include a baseline and multiple metrics. The boosted model's R² 0.490 sounds useful until persistence reaches 0.732 with lower error. The residual network reaches 0.867 and reduces MAE, but worst errors above 300 deaths remain. For planning, a calibrated interval and state-specific error monitoring are at least as important as the point estimate.")
    para(doc, "No causal claim can be made from these aggregate occurrence counts. Cluster differences could reflect age structure, population, coding practice, epidemic timing, and healthcare access. The COVID share of all deaths is not a population rate. The source is frozen and provisional, so a current deployment must revalidate the schema and retrain on current data.")
    subheading(doc, "Concrete recommendations")
    bullets(doc, [
        "Adopt a data contract that requires Group, Sex, Age Group, geography, period completeness, and suppression status in every analytical extract. Reject aggregations that mix totals with components.",
        "Store suppressed counts as interval-censored values with lower=1 and upper=9 fields. Use lower/midpoint/upper sensitivity scenarios for summaries that depend on small cells; never silently zero-fill.",
        "Use the four clusters as peer groups for review: older-age interventions for Cluster 1; respiratory-overlap and working-age male analysis for Cluster 3; timing-focused comparisons for Cluster 0; broad burden benchmarking for Cluster 2.",
        "Run the residual network as a challenger beside the lag-1 baseline. Suspend its operational use if rolling two-month WAPE exceeds the baseline by more than 5 percentage points, and retrain only after diagnosing drift.",
        "Add prediction intervals through blocked conformal calibration or quantile loss before staffing decisions. Publish MAE, RMSE, WAPE, and worst-jurisdiction errors each month.",
        "Replace the frozen extract with current CDC WONDER or another maintained official feed, then add population by age, vaccination, hospitalization, and variant indicators. Those variables are needed for rates and may improve turning-point forecasts.",
    ])
    para(doc, "The 5-percentage-point monitoring threshold is a governance recommendation, not a value estimated from this dataset. It creates a clear fallback rule while the baseline remains available. A production owner should revise it based on the cost of over- and underforecasting.")

    section(doc, 10, "Conclusion")
    para(doc, "The capstone delivered a complete data science pipeline on a difficult public-health table: 137,700 records were cleaned without erasing censoring; EDA quantified time, age, sex, geography, and respiratory association; standardized profile features produced four interpretable but modest clusters; chronological evaluation rejected a conventional boosted model; and a compact residual neural network improved future MAE to 34.253 and R² to 0.867.")
    para(doc, "The next stage is operational validation on a maintained CDC feed with population denominators, prediction intervals, and rolling drift checks. The recommended system keeps the previous-month forecast as a fallback and promotes the neural correction only while it continues to beat that baseline on future data.")
    section(doc, 11, "References"); references(doc, include_tf=True)
    return save_doc(doc, "Week_6_Integrative_Capstone_Project.docx")


def build_descriptions():
    descriptions = {
        1: "I used the CDC/NCHS Provisional COVID-19 Deaths by Sex and Age file because it is a realistic public dataset with 137,700 rows, three reporting grains, overlapping age categories, comma-formatted counts, and privacy suppression. I first loaded every field as text so pandas could not silently mix numeric and string types, then verified the 16-column schema, 54 geography labels, three sex levels, 17 age groups, and the monthly, yearly, and cumulative record structure. The most important finding during cleaning was that missing count values were not ordinary blanks: 39,430 COVID-19 cells and up to 44,864 pneumonia cells were suppressed under the NCHS 1–9 confidentiality rule, and every missing count appeared on a footnoted row. I therefore kept those values as unknown and added measure-specific suppression flags instead of filling them with zero, a mean, or a median. I also parsed dates and counts, used nullable Year and Month fields, checked duplicate grain keys, tested nonnegative and logical count relationships, and marked 5,508 rows from partial September 2023 or partial 2023 annual periods. The quality tests found zero duplicate keys, zero invalid date orders, and zero cases where COVID-19 exceeded total deaths. An IQR screen flagged California, Florida, and Texas as high-count jurisdictions, but I kept them because the values were plausible population-scale burdens rather than errors. The hardest parts were separating structural blanks from suppressed counts and preventing overlapping age bands from being summed. I handled both with explicit rules and a mutually exclusive age scheme for later work. This clean, auditable master table sets up Week 2 EDA without double-counting totals, treating censored deaths as zero, or comparing incomplete periods as though they were complete.",
        2: "I carried the cleaned CDC Provisional COVID-19 Deaths by Sex and Age data from Week 1 into an EDA focused on questions that can be answered at a defensible grain. I built separate analytical slices for the national monthly trend, cumulative age burden, male–female comparisons, jurisdiction death composition, and correlations among respiratory mortality series. The main safeguard was to filter one Group, Sex, and Age Group level before every aggregation; for age analysis I used 11 mutually exclusive bands rather than mixing the overlapping schemes in the source. The results were specific: the file contains 1,146,774 cumulative U.S. COVID-19 deaths, equal to 9.32% of 12,303,399 all-cause deaths in the covered period. January 2021 was the highest complete month with 105,565 COVID-19 deaths and a 28.25% share of all deaths. Ages 65 and older accounted for 868,831 deaths, or 75.76% of the cumulative total. Across the exclusive age bands, male deaths exceeded female deaths by a ratio of 1.218, with the largest proportional gap, 1.75, at ages 45–54. New Jersey had the highest COVID share of total deaths among the 52 non-overlapping jurisdictions at 11.40%, while Vermont was lowest at 4.10%. COVID-19 and pneumonia monthly deaths had a Spearman correlation of 0.919. I used annotated line, grouped-bar, dual-axis age, ranked jurisdiction, and correlation heatmap visuals, and explained each both technically and analytically. The main challenges were the partial September 2023 period, overlapping New York/New York City geography, and the risk of calling a death-share a population risk rate. I excluded or labeled each case explicitly. The work establishes scale-free, interpretable features for Week 3 clustering.",
        3: "I used the cleaned CDC mortality data to segment 52 non-overlapping jurisdictions by profile rather than raw size. The national total and New York City were excluded because the first is an aggregate and the second overlaps New York. I engineered seven percentages covering COVID-19 and pneumonia shares of all deaths, the pneumonia–COVID overlap, influenza share, the percentage of COVID-19 deaths among ages 65+, male share, and the percentage of 2020–2022 COVID-19 deaths that occurred in 2021. This design was important because clustering raw death counts would mainly separate high-population states from small ones. All 364 feature cells were observed, and I standardized the features before applying K-Means. I tested k=2 through k=8 with 50 initializations per candidate, then selected four clusters because k=4 had the highest silhouette score, 0.222; k=7 was close at 0.220, so I reported the modest separation rather than presenting the groups as fixed natural categories. The four clusters contained 8, 13, 23, and 8 jurisdictions. One profile had the lowest COVID share but high male and 2021 concentration; another had an 82.672% age-65+ share and the lowest pneumonia overlap; the largest cluster represented a broad higher-burden middle pattern; and the final cluster combined the highest pneumonia share and COVID–pneumonia overlap with a younger, more male-skewed profile. PCA retained 62.19% of variance in two dimensions, while a Ward dendrogram provided a second similarity view. The hardest decisions were preventing population size from dominating and choosing k when scores were close. Percentage features, scaling, repeated starts, and transparent metric reporting addressed those problems. These profiles create peer groups for public-health comparison and set up Week 4 forecasting without turning cluster membership into a causal label.",
        4: "I framed the cleaned CDC data as a one-month-ahead regression problem for 52 jurisdictions. The target was each jurisdiction's monthly COVID-19 death count, and every predictor was restricted to information from prior months: three COVID-19 lags, a previous three-month mean, lagged all-cause and pneumonia deaths, seasonal sine and cosine terms, a time index, and one-hot jurisdiction identity. I excluded partial September 2023 and never treated a suppressed target as an exact observed count. A midpoint of 5 was used only for a lagged predictor when the source disclosed that the true value lay between 1 and 9. I trained a regularized HistGradientBoostingRegressor on log1p counts and used expanding chronological validation rather than random splits, because random rows would leak pandemic-wave structure across train and test. The model had mixed validation behavior: it beat persistence in 2021 H2 and 2022 H1 but lost in 2022 H2. On the untouched January–August 2023 holdout, the model produced MAE 69.259, RMSE 116.878, R² 0.490, and WAPE 54.239%. The previous-month baseline was better on every measure, with MAE 48.892, RMSE 84.661, R² 0.732, and WAPE 38.289%. California in January 2023 was the largest model error: 1,327 actual versus 2,550.8 predicted. The main challenge was resisting the temptation to present positive R² as success when a simple baseline was stronger. I documented the failure directly and linked it to regime change and poor extrapolation during the 2023 decline. This evaluation sets up a more defensible Week 5 architecture that keeps the lag baseline as an explicit skip path and learns only a correction.",
        5: "I redesigned the Week 4 forecast as a compact TensorFlow residual neural network rather than asking a larger model to relearn the full death-count scale. The network predicts the difference between log1p current deaths and log1p previous-month deaths, so a zero output reproduces the strong persistence baseline. I used the same 61 leak-free inputs and chronological data partitions: 1,381 rows through June 2022 for fitting, 309 rows from July–December 2022 for validation, and 379 rows from January–August 2023 for testing. The architecture has 32 and 16 ReLU hidden units, 20% and 15% dropout, L2 regularization, and one linear correction output, for 2,529 trainable parameters. I chose ReLU because the standardized tabular inputs need non-saturating nonlinear interactions, a linear output because corrections can move in either direction, and Huber loss because wave shocks should not dominate updates as strongly as they would under squared loss. Early stopping restored epoch 4 after training stopped at epoch 24. The test results improved materially: MAE 34.253, RMSE 59.621, R² 0.867, and WAPE 26.825%, compared with baseline MAE 48.892 and WAPE 38.289%. That is a 29.94% MAE reduction. I also trained an unregularized reference for 160 epochs; its final training loss fell to 0.01497 while validation loss rose to 0.49378, providing direct evidence of overfitting. The hardest issues were the small sample, changing pandemic regimes, and rapid January–February transitions; Texas February 2023 still missed by 333.2 deaths. The work establishes a monitored challenger model for the capstone, while keeping the baseline available as a fallback.",
        6: "I completed an end-to-end data science pipeline using the CDC/NCHS Provisional COVID-19 Deaths by Sex and Age extract. I began with 137,700 raw rows and treated the source as a historical, frozen dataset through 23 September 2023. Cleaning preserved all records, parsed mixed count fields, added suppression and partial-period metadata, and verified zero duplicate grain keys, invalid dates, negative counts, or impossible total/overlap relationships. I did not zero-fill privacy-suppressed 1–9 counts, and I defined a mutually exclusive age scheme so aggregate analysis could not double-count deaths. The EDA quantified 1,146,774 cumulative U.S. COVID-19 deaths, a January 2021 peak of 105,565, a 75.76% age-65+ share, and a male/female ratio of 1.218. Seven standardized, scale-free features then produced four jurisdiction profiles with sizes 8, 13, 23, and 8; the silhouette score of 0.222 was useful but modest, so I treated the clusters as peer groups rather than fixed labels. For supervised forecasting, chronological validation exposed an important failure: a gradient-boosted model had positive R² but was worse than the previous-month baseline on every 2023 holdout metric. I responded by designing a 2,529-parameter residual neural network that learned a correction around that baseline. It reached MAE 34.253, RMSE 59.621, R² 0.867, and WAPE 26.825%, reducing baseline MAE by 29.94%. Training curves and an unregularized reference documented overfitting rather than discussing it abstractly. The main challenges were censoring, overlapping aggregates, weak cluster boundaries, temporal regime change, and limited deep-learning sample size. I recommend a current CDC feed, interval-aware suppression handling, prediction intervals, rolling WAPE monitoring, and automatic fallback to persistence when the challenger degrades. This capstone sets up operational validation with current data, population denominators, and external leading indicators.",
    }
    paths = []
    for week, text in descriptions.items():
        path = OUT / f"Week_{week}_SUBMISSION_DESCRIPTION.txt"
        path.write_text(
            f"SUBMISSION DESCRIPTION — WEEK {week}\n\n{text}\n",
            encoding="utf-8",
        )
        paths.append(path)
    return paths


if __name__ == "__main__":
    paths = [
        build_week1(), build_week2(), build_week3(),
        build_week4(), build_week5(), build_week6(),
        *build_descriptions(),
    ]
    for p in paths:
        print(p)
